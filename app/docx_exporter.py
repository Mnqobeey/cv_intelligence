from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import Table

from .constants import BASE_TEMPLATE_PATH, MASTER_TEMPLATE_PATH, LOGO_PATH
from .parsers import format_recruiter_date, normalize_recruiter_date_text, parse_experience_section, sanitize_entity_text
from .schema import CandidateProfileSchema, validate_export_payload

ACCENT = RGBColor(0, 0, 0)
ACCENT_DARK = RGBColor(0, 0, 0)
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(0, 0, 0)
LIGHT_FILL = "F2F2F2"
MID_FILL = "E6E6E6"
WHITE_FILL = "FFFFFF"
FONT_FAMILY = "Arial"
BODY_FONT_SIZE = 10
SUBHEADING_FONT_SIZE = 11
TOP_NAME_ROLE_FONT_SIZE = 18
TOP_META_FONT_SIZE = 11
SOLID_BLACK_BORDER = "000000"
SOLID_BLACK_BORDER_SIZE = "10"
CAREER_TABLE_MARKERS = {
    "{{CAREER_ROLE_COMPANY}}",
    "{{CAREER_ROLE_TITLE}}",
    "{{CAREER_ROLE_START}}",
    "{{CAREER_ROLE_END}}",
    "{{CLIENT_ENGAGEMENTS_BLOCK}}",
    "{{PROJECT_NAME}}",
    "{{PROJECT_DETAILS}}",
}
REQUIRED_TEMPLATE_MARKERS = {
    "{{FULL_NAME}}",
    "{{PROFESSIONAL_TITLE}}",
    "{{AVAILABILITY}}",
    "{{REGION}}",
    "{{CANDIDATE_SUMMARY}}",
    "{{SKILLS_BLOCK}}",
    "{{QUALIFICATION}}",
    "{{INSTITUTION}}",
    "{{END_DATE}}",
    "{{CERTIFICATION_1}}",
    "{{COMPANY}}",
    "{{POSITION}}",
    "{{START_DATE}}",
    "{{CAREER_ROLE_COMPANY}}",
    "{{CAREER_ROLE_TITLE}}",
    "{{CAREER_ROLE_START}}",
    "{{CAREER_ROLE_END}}",
    "{{CLIENT_ENGAGEMENTS_BLOCK}}",
    "{{PROJECT_NAME}}",
    "{{PROJECT_DETAILS}}",
}
ALLOWED_TEMPLATE_LITERALS = {
    "Experience Innovation",
    "Recruiter-ready professional profile",
    "Candidate Summary",
    "Skills",
    "Qualification",
    "Certifications",
    "Career Summary",
    "Career History",
    "Qualification",
    "Institution",
    "End Date",
    "Company",
    "Position",
    "Start Date",
    "Company Name",
    "Job Title",
    "Start date",
    "End date",
    "Availability:",
    "Availability",
    "Region:",
    "Region",
    "Responsibilities:",
    "Client Engagements:",
    "Projects (Linked to this role):",
    "Project",
    "Details",
    "CestaSoft • Experience Innovation • recruiter-ready professional profile",
}


@dataclass
class FinalProfilePayload:
    identity: Dict[str, str]
    summary: str
    skills: List[Dict[str, str]] = field(default_factory=list)
    qualifications: List[Dict[str, str]] = field(default_factory=list)
    certifications: List[Dict[str, str]] = field(default_factory=list)
    career_summary: List[Dict[str, str]] = field(default_factory=list)
    career_history: List[Dict[str, Any]] = field(default_factory=list)


def _clean(value: Optional[str]) -> str:
    value = sanitize_entity_text(value) or ""
    value = re.sub(r"\s+", " ", value).strip(" |,-")
    if value.lower() in {"", "n/a", "na", "none", "null", "not provided", "not specified"}:
        return ""
    return value


def _split_lines(text: str) -> List[str]:
    return [line.strip() for line in (text or "").splitlines() if _clean(line)]


def _split_structured_lines(text: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for raw in _split_lines(text):
        if raw.lower().startswith("summary:"):
            continue
        parts = [_clean(part) for part in re.split(r"\s*\|\s*", raw)]
        parts = [part for part in parts if part]
        if parts:
            rows.append(parts)
    return rows


def _dedupe_preserve_order(values: Iterable[str]) -> List[str]:
    seen = set()
    result: List[str] = []
    for value in values:
        item = _clean(value)
        if not item:
            continue
        key = item.casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def _compact_skill_rows(rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    compact: List[Dict[str, str]] = []
    core_items: List[str] = []
    seen_rows = set()
    seen_core = set()

    for row in rows:
        category = _clean(row.get("category"))
        items = _clean(row.get("items"))
        if not items:
            continue

        if category and items.lower().startswith(f"{category.lower()}:"):
            items = _clean(items.split(":", 1)[1])
        elif ":" in items:
            maybe_category, maybe_items = items.split(":", 1)
            if _clean(maybe_category) and _clean(maybe_items):
                category = _clean(maybe_category)
                items = _clean(maybe_items)

        if category in {"", "Core Skills"}:
            for item in [_clean(part) for part in re.split(r"\s*,\s*", items) if _clean(part)] or [items]:
                key = item.casefold()
                if key in seen_core:
                    continue
                seen_core.add(key)
                core_items.append(item)
            continue

        row_key = (category.casefold(), items.casefold())
        if row_key in seen_rows:
            continue
        seen_rows.add(row_key)
        compact.append({"category": category, "items": items})

    if core_items:
        compact.insert(0, {"category": "Core Skills", "items": ", ".join(core_items)})
    return compact


def _skill_display_lines(rows: List[Dict[str, str]]) -> List[str]:
    lines: List[str] = []
    for row in rows:
        category = _clean(row.get("category"))
        items = _clean(row.get("items"))
        if not items:
            continue
        if category in {"", "Core Skills"}:
            split_items = [_clean(part) for part in re.split(r"\s*,\s*", items) if _clean(part)]
            if split_items:
                lines.extend(split_items)
            else:
                lines.append(items)
            continue
        lines.append(f"{category}: {items}")
    return lines


def _career_detail_heading(entry: Dict[str, Any]) -> str:
    explicit = _clean(entry.get("detail_heading"))
    if explicit:
        return explicit.rstrip(":") + ":"

    company = _clean(entry.get("company")).lower()
    bullets = [_clean(item) for item in entry.get("responsibilities", []) if _clean(item)]
    client_style_count = sum(1 for bullet in bullets if re.match(r"^[A-Z][A-Za-z0-9&/().,' -]{1,60}:\s+\S", bullet))
    if "consulting partner" in company and client_style_count >= 2:
        return "Client Engagements:"
    return "Responsibilities:"


def _client_engagement_line(client: Dict[str, Any]) -> str:
    client_name = _clean(client.get("client_name"))
    project_name = _clean(client.get("project_name") or client.get("programme"))
    details = "; ".join(_clean(item) for item in client.get("responsibilities", []) if _clean(item))
    label = client_name or project_name
    if label and details:
        return f"{label}: {details}"
    return label or details


def _project_rows_from_clients(clients: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    seen = set()
    for client in clients:
        name = _clean(client.get("project_name") or client.get("programme"))
        client_name = _clean(client.get("client_name"))
        details = "; ".join(_clean(item) for item in client.get("responsibilities", []) if _clean(item))
        if client_name and client_name.lower() != name.lower():
            details = " | ".join(part for part in [client_name, details] if part)
        if not name:
            continue
        key = (name.casefold(), details.casefold())
        if key in seen:
            continue
        seen.add(key)
        rows.append({"name": name, "details": details})
    return rows


def _looks_like_client_engagement(line: str) -> bool:
    cleaned = _clean(line)
    return bool(re.match(r"^[A-Z][A-Za-z0-9&/().,' -]{1,60}:\s+\S", cleaned))


def _skills_from_state(state: Dict[str, str], profile: Optional[Dict[str, Any]]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for line in _split_lines(state.get("skills", "")):
        if ":" in line:
            category, items = line.split(":", 1)
            rows.append({"category": _clean(category), "items": _clean(items)})
        else:
            rows.append({"category": "Core Skills", "items": _clean(line)})
    if rows:
        return _compact_skill_rows(rows)
    declared = _dedupe_preserve_order(((profile or {}).get("skills") or {}).get("declared", []))
    return _compact_skill_rows([{"category": "Core Skills", "items": ", ".join(declared)}]) if declared else []


def _certifications_from_state(state: Dict[str, str], profile: Optional[Dict[str, Any]]) -> List[Dict[str, str]]:
    def parse_line(raw: str) -> Dict[str, str]:
        raw = _normalize_bullet_item(raw)
        parts = [_normalize_bullet_item(part) for part in re.split(r"\s*\|\s*", raw) if _normalize_bullet_item(part)]
        if len(parts) >= 3:
            return {"name": parts[0], "provider": parts[1], "year": parts[2]}
        if len(parts) == 2:
            if re.fullmatch(r"(?:19|20)\d{2}|Present|In Progress", parts[1], re.I):
                return {"name": parts[0], "provider": "", "year": parts[1]}
            return {"name": parts[0], "provider": parts[1], "year": ""}
        match = re.search(r"\((\d{4}|In Progress|Present)\)$", raw, re.I)
        year = match.group(1) if match else ""
        name = re.sub(r"\s*\((\d{4}|In Progress|Present)\)$", "", raw, flags=re.I).strip()
        return {"name": _clean(name), "provider": "", "year": _clean(year)}

    parsed: List[Dict[str, str]] = []
    seen = set()
    for raw in _split_lines(state.get("certifications", "")):
        item = parse_line(raw)
        key = (item["name"].casefold(), item["provider"].casefold(), item["year"].casefold())
        if not item["name"] or key in seen:
            continue
        seen.add(key)
        parsed.append(item)
    if parsed:
        return parsed
    fallback: List[Dict[str, str]] = []
    for raw in _dedupe_preserve_order((profile or {}).get("certifications", [])):
        item = parse_line(raw)
        if item["name"]:
            fallback.append(item)
    return fallback


def _parse_date_parts(dates: str) -> tuple[str, str]:
    separators = f"-{chr(8211)}{chr(8212)}"
    parts = re.split(rf"\s+(?:[{separators}]|to)\s+", _clean(dates), maxsplit=1, flags=re.I)
    if len(parts) == 2:
        return format_recruiter_date(parts[0]), format_recruiter_date(parts[1])
    return format_recruiter_date(_clean(dates)), ""


def build_final_profile_payload(state: Dict[str, str], profile: Optional[Dict[str, Any]] = None) -> FinalProfilePayload:
    career_history = _career_history_from_state(state, profile)
    availability_value = sanitize_entity_text(state.get("availability")) or ""
    region_value = sanitize_entity_text(state.get("region")) or ""
    identity = {
        "full_name": _clean(state.get("full_name")) or "Candidate Name",
        "headline": _clean(state.get("headline")) or "Professional Profile",
        "availability": availability_value if availability_value.lower() == "not provided" else _clean(availability_value),
        "region": region_value if region_value.lower() == "not provided" else _clean(region_value),
        "email": _clean(state.get("email")),
        "phone": _clean(state.get("phone")),
        "location": _clean(state.get("location")),
        "linkedin": _clean(state.get("linkedin")),
        "portfolio": _clean(state.get("portfolio")),
    }
    return FinalProfilePayload(
        identity=identity,
        summary=_clean(state.get("summary")) or _clean((profile or {}).get("summary")),
        skills=_skills_from_state(state, profile),
        qualifications=_qualifications_from_state(state, profile),
        certifications=_certifications_from_state(state, profile),
        career_summary=_career_summary_from_history(career_history),
        career_history=career_history,
    )


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def _set_paragraph_text(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float = 10, color: RGBColor | None = None) -> None:
    paragraph.clear()
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size)
    if color is None:
        color = BLACK
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(key), FONT_FAMILY)
    if color:
        run.font.color.rgb = color


def _normalize_bullet_item(text: str) -> str:
    cleaned = _clean(text)
    cleaned = re.sub(r"^(?:[\u2022\u00b7\-\*\u25cf\?]+\s*)+", "", cleaned)
    return _clean(cleaned)


def _placeholder_style(marker: str) -> Dict[str, Any]:
    styles: Dict[str, Dict[str, Any]] = {
        "{{FULL_NAME}}": {"bold": True, "size": TOP_NAME_ROLE_FONT_SIZE, "color": BLACK},
        "{{HEADLINE}}": {"bold": True, "size": TOP_NAME_ROLE_FONT_SIZE, "color": BLACK},
        "{{PROFESSIONAL_TITLE}}": {"bold": True, "size": TOP_NAME_ROLE_FONT_SIZE, "color": BLACK},
        "{{AVAILABILITY_LINE}}": {"bold": True, "size": TOP_META_FONT_SIZE, "color": BLACK},
        "{{REGION_LINE}}": {"bold": True, "size": TOP_META_FONT_SIZE, "color": BLACK},
        "{{AVAILABILITY}}": {"bold": True, "size": TOP_META_FONT_SIZE, "color": BLACK},
        "{{REGION}}": {"bold": True, "size": TOP_META_FONT_SIZE, "color": BLACK},
        "{{LOCATION_LINE}}": {"size": BODY_FONT_SIZE, "color": ACCENT_DARK},
        "{{CANDIDATE_SUMMARY}}": {"size": BODY_FONT_SIZE},
        "{{SKILLS_BLOCK}}": {"size": BODY_FONT_SIZE},
        "{{CERTIFICATION_ITEM}}": {"size": BODY_FONT_SIZE},
    }
    return styles.get(marker, {"size": BODY_FONT_SIZE})


def _replace_placeholder_text(paragraph, marker: str, value: str) -> None:
    if marker not in paragraph.text:
        return
    text = paragraph.text.replace(marker, value)
    style = _placeholder_style(marker)
    _set_paragraph_text(paragraph, text, **style)
    if marker in {"{{AVAILABILITY}}", "{{REGION}}", "{{AVAILABILITY_LINE}}", "{{REGION_LINE}}"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if marker == "{{CANDIDATE_SUMMARY}}":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _clear_paragraph_list_format(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)
    if paragraph.style and paragraph.style.name == "List Paragraph":
        paragraph.style = paragraph.part.document.styles["Normal"]


def _delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def _delete_table(table) -> None:
    tbl = table._element
    tbl.getparent().remove(tbl)


def _iter_document_text(document: Document) -> Iterable[str]:
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            yield text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = (paragraph.text or "").strip()
                    if text:
                        yield text


def _is_placeholder_only_text(text: str) -> bool:
    if "{{" not in text or "}}" not in text:
        return False
    stripped = re.sub(r"{{[A-Z0-9_]+}}", "", text)
    stripped = re.sub(r"\s+", " ", stripped).strip(" |-:•·")
    return stripped == "" or stripped in ALLOWED_TEMPLATE_LITERALS


def _template_has_only_allowed_text(document: Document) -> bool:
    for text in _iter_document_text(document):
        if text in ALLOWED_TEMPLATE_LITERALS or _is_placeholder_only_text(text):
            continue
        return False
    return True


def _template_contains_required_markers(document: Document) -> bool:
    full_text = "\n".join(_iter_document_text(document))
    return all(marker in full_text for marker in REQUIRED_TEMPLATE_MARKERS)


def _template_has_expected_typography(document: Document) -> bool:
    normal = document.styles["Normal"].font.name or ""
    if normal and normal != FONT_FAMILY:
        return False
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text in {"Candidate Summary", "Skills", "Qualification", "Certifications", "Career Summary", "Career History"}:
            if not paragraph.runs:
                return False
            if (paragraph.runs[0].font.name or "") != FONT_FAMILY:
                return False
    return True


def _find_paragraph(document: Document, marker: str):
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            return paragraph
    return None


def _find_table_with_marker(document: Document, marker: str):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if marker in cell.text:
                    return table
    return None


def _insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def _apply_cell_borders(cell, *, color: str = "D0D7DE", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _append_cell_paragraph(cell, text: str, *, bold: bool = False, italic: bool = False, size: float = 10, color: RGBColor | None = None, left_indent: float = 0.0, first_line_indent: float = 0.0, space_after: float = 1.2):
    paragraph = cell.add_paragraph()
    _set_paragraph_text(paragraph, text, bold=bold, italic=italic, size=size, color=color)
    if left_indent:
        paragraph.paragraph_format.left_indent = Cm(left_indent)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def _style_table_cells(table, *, color: str = SOLID_BLACK_BORDER, size: str = SOLID_BLACK_BORDER_SIZE) -> None:
    for row in table.rows:
        for cell in row.cells:
            _apply_cell_borders(cell, color=color, size=size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _align_table_paragraphs(table, alignment) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment


def _style_qualifications_table(table) -> None:
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, MID_FILL)
        if cell.paragraphs and cell.text.strip():
            _set_paragraph_text(
                cell.paragraphs[0],
                cell.text.strip(),
                bold=True,
                color=BLACK,
            )
    for row in table.rows[1:]:
        for cell in row.cells:
            _set_cell_shading(cell, WHITE_FILL)
            if cell.paragraphs and cell.text.strip():
                _set_paragraph_text(cell.paragraphs[0], cell.text.strip(), size=10, color=BLACK)
    _style_table_cells(table)
    _align_table_paragraphs(table, WD_ALIGN_PARAGRAPH.CENTER)


def _style_career_history_table(table) -> None:
    _style_table_cells(table)


def _add_career_history_template_table(doc: Document):
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.4), Inches(2.9), Inches(1.0), Inches(1.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    company_cells = table.rows[0].cells
    company_cells[1].merge(company_cells[3])
    _set_cell_shading(company_cells[0], MID_FILL)
    _set_paragraph_text(company_cells[0].paragraphs[0], "Company Name", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(company_cells[1].paragraphs[0], "{{CAREER_ROLE_COMPANY}}", size=10, color=BLACK)

    title_cells = table.rows[1].cells
    _set_cell_shading(title_cells[0], LIGHT_FILL)
    _set_cell_shading(title_cells[2], LIGHT_FILL)
    _set_cell_shading(title_cells[3], LIGHT_FILL)
    _set_paragraph_text(title_cells[0].paragraphs[0], "Job Title", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(title_cells[1].paragraphs[0], "{{CAREER_ROLE_TITLE}}", size=10)
    _set_paragraph_text(title_cells[2].paragraphs[0], "Start date", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(title_cells[3].paragraphs[0], "End date", bold=True, size=10, color=ACCENT_DARK)

    date_cells = table.rows[2].cells
    date_cells[0].merge(date_cells[1])
    _set_paragraph_text(date_cells[0].paragraphs[0], "", size=8)
    _set_paragraph_text(date_cells[2].paragraphs[0], "{{CAREER_ROLE_START}}", size=10)
    _set_paragraph_text(date_cells[3].paragraphs[0], "{{CAREER_ROLE_END}}", size=10)

    resp_cells = table.rows[3].cells
    resp_cell = resp_cells[0].merge(resp_cells[3])
    _set_cell_shading(resp_cell, "FBFCFE")
    _set_paragraph_text(resp_cell.paragraphs[0], "Responsibilities:", bold=True, size=10, color=ACCENT_DARK)
    _append_cell_paragraph(resp_cell, "{{RESPONSIBILITY_1}}", size=10, left_indent=0.45, first_line_indent=-0.2)
    _append_cell_paragraph(resp_cell, "{{RESPONSIBILITY_2}}", size=10, left_indent=0.45, first_line_indent=-0.2)
    spacer = resp_cell.add_paragraph()
    _set_paragraph_text(spacer, "", size=6)
    _set_paragraph_text(resp_cell.add_paragraph(), "Client Engagements:", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(resp_cell.add_paragraph(), "{{CLIENT_ENGAGEMENTS_BLOCK}}", size=10)
    spacer = resp_cell.add_paragraph()
    _set_paragraph_text(spacer, "", size=6)
    _set_paragraph_text(resp_cell.add_paragraph(), "Projects (Linked to this role):", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(resp_cell.add_paragraph(), "Project", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(resp_cell.add_paragraph(), "{{PROJECT_NAME}}", size=10)
    spacer = resp_cell.add_paragraph()
    _set_paragraph_text(spacer, "", size=6)
    _set_paragraph_text(resp_cell.add_paragraph(), "Details", bold=True, size=10, color=ACCENT_DARK)
    _set_paragraph_text(resp_cell.add_paragraph(), "{{PROJECT_DETAILS}}", size=10)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _apply_cell_borders(cell, color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
    return table


def _retarget_existing_template(doc: Document) -> None:
    from docx.text.paragraph import Paragraph

    def ensure_location_after(region_paragraph):
        if any("{{LOCATION_LINE}}" in p.text for p in doc.paragraphs):
            return
        new_p = OxmlElement("w:p")
        region_paragraph._p.addnext(new_p)
        para = Paragraph(new_p, region_paragraph._parent)
        _set_paragraph_text(para, "{{LOCATION_LINE}}", size=BODY_FONT_SIZE, color=ACCENT_DARK)

    heading_done = False
    cert_heading_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        lowered = text.lower()
        if not heading_done and text and " - " in text and not lowered.startswith(("availability", "region", "candidate summary", "skills")):
            _set_paragraph_text(paragraph, "{{FULL_NAME}} - {{HEADLINE}}", bold=True, size=TOP_NAME_ROLE_FONT_SIZE, color=BLACK)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)
            heading_done = True
            continue
        if text == "Availability:":
            _set_paragraph_text(paragraph, "{{AVAILABILITY_LINE}}", bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
        elif text.startswith("Region:"):
            _set_paragraph_text(paragraph, "{{REGION_LINE}}", bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
            ensure_location_after(paragraph)
        elif text == "Candidate Summary":
            _set_paragraph_text(paragraph, "Candidate Summary", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)
        elif text == "Skills":
            _set_paragraph_text(paragraph, "Skills", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)
        elif text == "Qualification":
            _set_paragraph_text(paragraph, "Qualification", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)
        elif text == "Certifications":
            _set_paragraph_text(paragraph, "Certifications", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)
            cert_heading_index = idx
        elif text == "Career Summary":
            _set_paragraph_text(paragraph, "Career Summary", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)
        elif text == "Career History":
            _set_paragraph_text(paragraph, "Career History", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)

    if cert_heading_index is not None:
        paras = doc.paragraphs
        cert_items = []
        scan_index = cert_heading_index + 1
        while scan_index < len(paras):
            candidate = paras[scan_index]
            value = candidate.text.strip()
            if not value:
                scan_index += 1
                continue
            if value in {"Career Summary", "Career History", "Skills", "Qualification"}:
                break
            cert_items.append(candidate)
            scan_index += 1
        if cert_items:
            _set_paragraph_text(cert_items[0], "{{CERTIFICATION_ITEM}}", size=BODY_FONT_SIZE)
            for extra in cert_items[1:]:
                _delete_paragraph(extra)
        if not any((p.text or '').strip() == "Certifications" for p in doc.paragraphs):
            anchor = cert_items[0] if cert_items else doc.paragraphs[cert_heading_index]
            heading = anchor.insert_paragraph_before() if hasattr(anchor, 'insert_paragraph_before') else None
            if heading is None:
                new_p = OxmlElement("w:p")
                anchor._p.addprevious(new_p)
                heading = Paragraph(new_p, anchor._parent)
            _set_paragraph_text(heading, "Certifications", bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)

    if doc.tables:
        summary_table = doc.tables[0]
        _set_paragraph_text(summary_table.cell(0, 0).paragraphs[0], "{{CANDIDATE_SUMMARY}}", size=BODY_FONT_SIZE)
        summary_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    if len(doc.tables) > 1:
        skills_table = doc.tables[1]
        _set_paragraph_text(skills_table.cell(0, 0).paragraphs[0], "{{SKILLS_BLOCK}}", size=BODY_FONT_SIZE)
        if len(skills_table.columns) > 1:
            skills_table.cell(0, 1).text = ""

    if doc.tables:
        qual_table = next((t for t in doc.tables if any("Qualification" in c.text for r in t.rows for c in r.cells)), None)
        if qual_table and len(qual_table.rows) >= 2:
            row = qual_table.rows[1].cells
            _set_paragraph_text(row[0].paragraphs[0], "{{QUALIFICATION_DEGREE}}", size=BODY_FONT_SIZE)
            _set_paragraph_text(row[1].paragraphs[0], "{{QUALIFICATION_INSTITUTION}}", size=BODY_FONT_SIZE)
            _set_paragraph_text(row[2].paragraphs[0], "{{QUALIFICATION_YEAR}}", size=BODY_FONT_SIZE)
            while len(qual_table.rows) > 2:
                qual_table._tbl.remove(qual_table.rows[-1]._tr)

        summary_tables = [t for t in doc.tables if any("Company" in c.text for r in t.rows for c in r.cells) and any("Position" in c.text for r in t.rows for c in r.cells) and any("Start Date" in c.text for r in t.rows for c in r.cells)]
        if summary_tables:
            cs_table = summary_tables[0]
            if len(cs_table.rows) >= 2:
                row = cs_table.rows[1].cells
                _set_paragraph_text(row[0].paragraphs[0], "{{CAREER_SUMMARY_COMPANY}}", size=BODY_FONT_SIZE)
                _set_paragraph_text(row[1].paragraphs[0], "{{CAREER_SUMMARY_POSITION}}", size=BODY_FONT_SIZE)
                _set_paragraph_text(row[2].paragraphs[0], "{{CAREER_SUMMARY_START}}", size=BODY_FONT_SIZE)
                _set_paragraph_text(row[3].paragraphs[0], "{{CAREER_SUMMARY_END}}", size=BODY_FONT_SIZE)
                while len(cs_table.rows) > 2:
                    cs_table._tbl.remove(cs_table.rows[-1]._tr)

    history_tables = [t for t in doc.tables if any("Company Name" in c.text for r in t.rows for c in r.cells) and any("Job Title" in c.text for r in t.rows for c in r.cells)]
    if history_tables:
        for table in history_tables:
            _delete_table(table)
        marker = _find_paragraph(doc, "Career History")
        if marker:
            item = _find_paragraph(doc, "{{CAREER_HISTORY_ITEM}}") or _insert_paragraph_after(marker, "{{CAREER_HISTORY_ITEM}}")
            item.paragraph_format.space_after = Pt(3)
            _add_career_history_template_table(doc)

def _create_master_template(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.45)
    sec.bottom_margin = Cm(1.35)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_FAMILY
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.font.color.rgb = BLACK

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(1.8)
    header.columns[1].width = Inches(4.9)
    if LOGO_PATH.exists():
        run = header.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.55))
    p = header.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_text(p, "Experience Innovation", bold=True, size=11, color=ACCENT)
    p2 = header.cell(0, 1).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_text(p2, "Recruiter-ready professional profile", size=9, color=MUTED)

    divider = doc.add_paragraph()
    pPr = divider._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E57E25")
    pbdr.append(bottom)
    pPr.append(pbdr)

    p = doc.add_paragraph("{{FULL_NAME}}")
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = FONT_FAMILY
    p.runs[0].font.size = Pt(TOP_NAME_ROLE_FONT_SIZE)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph("{{HEADLINE}}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = FONT_FAMILY
    p.runs[0].font.size = Pt(TOP_NAME_ROLE_FONT_SIZE)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(6)

    meta = doc.add_table(rows=1, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Inches(3.1)
    meta.columns[1].width = Inches(3.1)
    for cell, marker in zip(meta.rows[0].cells, ["{{AVAILABILITY_LINE}}", "{{REGION_LINE}}"]):
        _set_cell_shading(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_paragraph_text(cell.paragraphs[0], marker, bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
        _apply_cell_borders(cell, color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)

    for heading, marker in [
        ("Candidate Summary", "{{CANDIDATE_SUMMARY}}"),
        ("Skills", "{{SKILLS_BLOCK}}"),
        ("Qualification", "QUALIFICATIONS_TABLE"),
        ("Certifications", "{{CERTIFICATION_ITEM}}"),
        ("Career Summary", "CAREER_SUMMARY_TABLE"),
        ("Career History", "{{CAREER_HISTORY_ITEM}}"),
    ]:
        h = doc.add_paragraph(heading)
        h.runs[0].bold = True
        h.runs[0].font.name = FONT_FAMILY
        h.runs[0].font.size = Pt(SUBHEADING_FONT_SIZE)
        h.runs[0].font.color.rgb = BLACK
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        if heading == "Qualification":
            table = doc.add_table(rows=2, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for cell, txt in zip(hdr, ["Qualification", "Institution", "End Date"]):
                _set_cell_shading(cell, MID_FILL)
                _set_paragraph_text(cell.paragraphs[0], txt, bold=True, color=ACCENT_DARK)
                _apply_cell_borders(cell, color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
            row = table.rows[1].cells
            row[0].text = "{{QUALIFICATION_DEGREE}}"
            row[1].text = "{{QUALIFICATION_INSTITUTION}}"
            row[2].text = "{{QUALIFICATION_YEAR}}"
            _set_cell_shading(row[0], WHITE_FILL)
            _set_cell_shading(row[1], WHITE_FILL)
            _set_cell_shading(row[2], WHITE_FILL)
            for cell in row:
                _apply_cell_borders(cell, color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
        elif heading == "Career Summary":
            table = doc.add_table(rows=2, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for cell, txt in zip(hdr, ["Company", "Position", "Start Date", "End Date"]):
                _set_cell_shading(cell, MID_FILL)
                _set_paragraph_text(cell.paragraphs[0], txt, bold=True, color=ACCENT_DARK)
                _apply_cell_borders(cell, color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
            row = table.rows[1].cells
            row[0].text = "{{CAREER_SUMMARY_COMPANY}}"
            row[1].text = "{{CAREER_SUMMARY_POSITION}}"
            row[2].text = "{{CAREER_SUMMARY_START}}"
            row[3].text = "{{CAREER_SUMMARY_END}}"
            for cell in row:
                _apply_cell_borders(cell, color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
        elif heading == "Career History":
            p = doc.add_paragraph(marker)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            _add_career_history_template_table(doc)
        else:
            p = doc.add_paragraph(marker)
            p.paragraph_format.space_after = Pt(3)
            if heading == "Career Summary":
                p.runs[0].font.size = Pt(BODY_FONT_SIZE)
            else:
                p.runs[0].font.size = Pt(BODY_FONT_SIZE)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_text(footer, "CestaSoft • Experience Innovation • recruiter-ready professional profile", size=8.5, color=MUTED)
    doc.save(str(path))

def _template_is_current(path: Path) -> bool:
    if not path.exists():
        return False
    document = Document(str(path))
    return (
        _template_contains_required_markers(document)
        and _template_has_only_allowed_text(document)
        and _template_has_expected_typography(document)
    )


def _ensure_heading_before_marker(document: Document, heading_text: str, marker: str) -> None:
    if any((paragraph.text or '').strip() == heading_text for paragraph in document.paragraphs):
        return
    marker_para = _find_paragraph(document, marker)
    if marker_para is None:
        return
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement("w:p")
    marker_para._p.addprevious(new_p)
    heading = Paragraph(new_p, marker_para._parent)
    _set_paragraph_text(heading, heading_text, bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)


def _set_section_headings_black(document: Document) -> None:
    section_labels = {
        "Candidate Summary",
        "Skills",
        "Qualification",
        "Certifications",
        "Career Summary",
        "Career History",
    }
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text in section_labels:
            _set_paragraph_text(paragraph, text, bold=True, size=SUBHEADING_FONT_SIZE, color=BLACK)


def _ensure_master_template() -> Path:
    path = MASTER_TEMPLATE_PATH if MASTER_TEMPLATE_PATH.exists() else BASE_TEMPLATE_PATH
    if path.exists():
        return path
    _create_master_template(path)
    return path


def _insert_paragraph_after_table(table, text: str = ""):
    new_p = OxmlElement("w:p")
    table._element.addnext(new_p)
    paragraph = table._parent.add_paragraph()
    paragraph._p = new_p
    if text:
        paragraph.add_run(text)
    return paragraph


def _insert_table_after_table(table, template_source) -> Table:
    new_tbl = deepcopy(template_source._tbl if hasattr(template_source, "_tbl") else template_source)
    table._element.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def _insert_table_after_paragraph(paragraph, template_source) -> Table:
    new_tbl = deepcopy(template_source._tbl if hasattr(template_source, "_tbl") else template_source)
    paragraph._p.addnext(new_tbl)
    return Table(new_tbl, paragraph._parent)


def _add_role_detail_bullet(cell, text: str) -> None:
    paragraph = cell.add_paragraph()
    _set_paragraph_text(paragraph, f"\u2022 {text}", size=10)
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.first_line_indent = Cm(-0.2)
    paragraph.paragraph_format.space_after = Pt(1.2)


def _fill_career_history_table(table, entry: Dict[str, Any]) -> None:
    mapping = {
        "{{CAREER_ROLE_COMPANY}}": _clean(entry.get("company")),
        "{{CAREER_ROLE_TITLE}}": _clean(entry.get("position")),
        "{{CAREER_ROLE_START}}": _clean(entry.get("start_date")) or _parse_date_parts(entry.get("dates", ""))[0],
        "{{CAREER_ROLE_END}}": _clean(entry.get("end_date")) or _parse_date_parts(entry.get("dates", ""))[1],
    }
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for marker, value in mapping.items():
                    if marker in paragraph.text:
                        _set_paragraph_text(paragraph, paragraph.text.replace(marker, value), size=10)
                        if marker == "{{CAREER_ROLE_COMPANY}}":
                            paragraph.runs[0].font.color.rgb = BLACK
                            paragraph.runs[0].bold = False

    if len(table.rows) < 4:
        _style_career_history_table(table)
        return
    detail_cell = table.rows[3].cells[0]
    for extra in detail_cell.paragraphs[1:]:
        _delete_paragraph(extra)
    first_para = detail_cell.paragraphs[0]
    _set_paragraph_text(first_para, "", size=4)

    responsibilities = [_clean(b) for b in entry.get("responsibilities", []) if _clean(b)]
    client_engagements = [_clean(item) for item in entry.get("client_engagements", []) if _clean(item)]
    projects = [project for project in entry.get("projects", []) if _clean((project or {}).get("name"))]
    if not responsibilities and _clean(entry.get("summary")):
        responsibilities = [_clean(entry.get("summary"))]

    sections_added = False

    def add_section(title: str, lines: List[str]) -> None:
        nonlocal sections_added
        if not lines:
            return
        heading = detail_cell.add_paragraph() if sections_added else first_para
        _set_paragraph_text(heading, title, bold=True, size=10, color=ACCENT_DARK)
        heading.paragraph_format.space_after = Pt(1.5)
        for line in lines[:8]:
            _add_role_detail_bullet(detail_cell, line)
        sections_added = True

    add_section("Responsibilities:", responsibilities)
    add_section("Client Engagements:", client_engagements)

    if projects:
        project_heading = detail_cell.add_paragraph() if sections_added else first_para
        _set_paragraph_text(project_heading, "Projects (Linked to this role):", bold=True, size=10, color=ACCENT_DARK)
        project_heading.paragraph_format.space_after = Pt(1.5)
        for project in projects[:4]:
            label = detail_cell.add_paragraph()
            _set_paragraph_text(label, "Project", bold=True, size=10, color=ACCENT_DARK)
            label.paragraph_format.space_after = Pt(0.6)
            value = detail_cell.add_paragraph()
            _set_paragraph_text(value, _clean(project.get("name")), size=10)
            value.paragraph_format.space_after = Pt(0.6)
            details = _clean(project.get("details"))
            if details:
                details_label = detail_cell.add_paragraph()
                _set_paragraph_text(details_label, "Details", bold=True, size=10, color=ACCENT_DARK)
                details_label.paragraph_format.space_after = Pt(0.6)
                details_para = detail_cell.add_paragraph()
                _set_paragraph_text(details_para, details, size=10)
                details_para.paragraph_format.space_after = Pt(1.0)

    if not sections_added and not projects:
        _delete_table(table)
        return
    _style_career_history_table(table)
    return
    for row in table.rows:
        for cell in row.cells:
            if any("{{CAREER_ROLE_BULLETS}}" in p.text for p in cell.paragraphs):
                bullet_cell = cell
                break
        if bullet_cell:
            break
    if bullet_cell is None:
        return
    first_para = bullet_cell.paragraphs[0]
    _set_paragraph_text(first_para, _career_detail_heading(entry), bold=True, size=9.4, color=ACCENT_DARK)
    first_para.paragraph_format.space_after = Pt(1.5)
    bullets = [_clean(b) for b in entry.get("responsibilities", []) if _clean(b)]
    if not bullets and _clean(entry.get("summary")):
        bullets = [_clean(entry.get("summary"))]
    for extra in bullet_cell.paragraphs[1:]:
        _delete_paragraph(extra)
    for bullet in bullets[:8]:
        paragraph = bullet_cell.add_paragraph()
        _set_paragraph_text(paragraph, f"• {bullet}", size=9.4)
        paragraph.paragraph_format.left_indent = Cm(0.45)
        paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        paragraph.paragraph_format.space_after = Pt(1.2)



def _remove_heading_for_marker(document: Document, marker: str) -> None:
    para = _find_paragraph(document, marker)
    if not para:
        return
    paras = list(document.paragraphs)
    prev = None
    for idx, item in enumerate(paras):
        if item._p is para._p:
            if idx > 0:
                prev = paras[idx - 1]
            break
    if prev and prev.text.strip() in {"Candidate Summary", "Skills", "Qualification", "Certifications", "Career Summary", "Career History"}:
        _delete_paragraph(prev)
    _delete_paragraph(para)


def _populate_simple_placeholder(document: Document, marker: str, value: str) -> None:
    for paragraph in document.paragraphs:
        _replace_placeholder_text(paragraph, marker, value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_placeholder_text(paragraph, marker, value)


def _populate_meta_lines(document: Document, identity: Dict[str, str]) -> None:
    raw_availability = sanitize_entity_text(identity.get("availability")) or ""
    raw_region = sanitize_entity_text(identity.get("region")) or ""
    raw_location = sanitize_entity_text(identity.get("location")) or ""
    availability = raw_availability if raw_availability.lower() == "not provided" else _clean(raw_availability)
    region = raw_region if raw_region.lower() == "not provided" else _clean(raw_region)
    location = raw_location if raw_location.lower() == "not provided" else _clean(raw_location)
    availability_line = f"Availability: {availability}" if availability else ""
    region_line = f"Region: {region}" if region else ""
    location_line = f"Location: {location}" if location else ""
    _populate_simple_placeholder(document, "{{AVAILABILITY_LINE}}", availability_line)
    _populate_simple_placeholder(document, "{{REGION_LINE}}", region_line)
    _populate_simple_placeholder(document, "{{AVAILABILITY}}", availability)
    _populate_simple_placeholder(document, "{{REGION}}", region)
    _populate_simple_placeholder(document, "{{LOCATION_LINE}}", location_line)

    table = None
    for t in document.tables:
        if any("Availability:" in c.text or "Region:" in c.text or "{{AVAILABILITY_LINE}}" in c.text for r in t.rows for c in r.cells):
            table = t
            break

    existing_texts = {paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()}
    headline_paragraph = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip() in {"{{HEADLINE}}", identity.get("headline", "").strip()}), None)

    if table is not None and (availability_line or region_line):
        anchor = headline_paragraph
        if availability_line and availability_line not in existing_texts and anchor is not None:
            anchor = _insert_paragraph_after(anchor)
            anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_text(anchor, availability_line, bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
            anchor.paragraph_format.space_after = Pt(2)
        if region_line and region_line not in existing_texts and anchor is not None:
            anchor = _insert_paragraph_after(anchor)
            anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_text(anchor, region_line, bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
            anchor.paragraph_format.space_after = Pt(6)
        if location_line and location_line not in existing_texts and anchor is not None:
            anchor = _insert_paragraph_after(anchor)
            anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_text(anchor, location_line, size=BODY_FONT_SIZE, color=ACCENT_DARK)
            anchor.paragraph_format.space_after = Pt(6)
        _delete_table(table)

    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text == "":
            continue
        if availability and text == availability_line:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_text(paragraph, text, bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
        elif region and text == region_line:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_text(paragraph, text, bold=True, size=TOP_META_FONT_SIZE, color=BLACK)
        elif location and text == location_line:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_text(paragraph, text, size=BODY_FONT_SIZE, color=ACCENT_DARK)
        elif text in {"Location:", "{{LOCATION_LINE}}"} and not location:
            _delete_paragraph(paragraph)
        elif text == "Availability:" and not availability:
            _delete_paragraph(paragraph)
        elif text == "Region:" and not region:
            _delete_paragraph(paragraph)


def _populate_list_block(document: Document, marker: str, items: List[str], *, bullet: bool = False) -> None:
    items = [_normalize_bullet_item(item) for item in items if _normalize_bullet_item(item)]
    paragraph = _find_paragraph(document, marker)
    if paragraph:
        if not items:
            _remove_heading_for_marker(document, marker)
            return
        anchor = paragraph
        for item in items:
            anchor = _insert_paragraph_after(anchor)
            _set_paragraph_text(anchor, ("\u2022 " if bullet else "") + item, size=BODY_FONT_SIZE)
            if bullet:
                anchor.paragraph_format.left_indent = Cm(0.4)
                anchor.paragraph_format.first_line_indent = Cm(-0.2)
        _delete_paragraph(paragraph)
        return

    table = _find_table_with_marker(document, marker)
    if not table:
        return
    if not items:
        _delete_table(table)
        return
    if marker == "{{SKILLS_BLOCK}}":
        skills_heading = next((para for para in document.paragraphs if para.text.strip() == "Skills"), None)
        if skills_heading is not None:
            anchor = skills_heading
            for item in items:
                anchor = _insert_paragraph_after(anchor)
                _set_paragraph_text(anchor, ("\u2022 " if bullet else "") + item, size=BODY_FONT_SIZE)
                if bullet:
                    anchor.paragraph_format.left_indent = Cm(0.4)
                    anchor.paragraph_format.first_line_indent = Cm(-0.2)
            _delete_table(table)
            return
    target_cell = None
    for row in table.rows:
        for cell in row.cells:
            if marker in cell.text:
                target_cell = cell
                break
        if target_cell:
            break
    if target_cell is None:
        return
    for paragraph in list(target_cell.paragraphs):
        if paragraph is target_cell.paragraphs[0]:
            _set_paragraph_text(paragraph, items[0], size=BODY_FONT_SIZE)
            if bullet:
                _set_paragraph_text(paragraph, f"\u2022 {items[0]}", size=BODY_FONT_SIZE)
                paragraph.paragraph_format.left_indent = Cm(0.4)
                paragraph.paragraph_format.first_line_indent = Cm(-0.2)
            continue
        _delete_paragraph(paragraph)
    for item in items[1:]:
        para = target_cell.add_paragraph()
        _set_paragraph_text(para, ("\u2022 " if bullet else "") + item, size=BODY_FONT_SIZE)
        if bullet:
            para.paragraph_format.left_indent = Cm(0.4)
            para.paragraph_format.first_line_indent = Cm(-0.2)


def _populate_certifications_block(document: Document, items: List[str]) -> None:
    items = [_normalize_bullet_item(item) for item in items if _normalize_bullet_item(item)]
    markers = [f"{{{{CERTIFICATION_{idx}}}}}" for idx in range(1, 4)]
    paragraphs = [paragraph for paragraph in document.paragraphs if any(marker in paragraph.text for marker in markers)]
    if not paragraphs:
        _populate_list_block(document, "{{CERTIFICATION_ITEM}}", items, bullet=True)
        return
    if not items:
        for paragraph in paragraphs:
            _delete_paragraph(paragraph)
        for paragraph in list(document.paragraphs):
            if paragraph.text.strip() == "Certifications":
                _delete_paragraph(paragraph)
                break
        return
    for index, paragraph in enumerate(paragraphs):
        if index < len(items):
            _clear_paragraph_list_format(paragraph)
            _set_paragraph_text(paragraph, f"\u2022 {items[index]}", size=BODY_FONT_SIZE)
            paragraph.paragraph_format.left_indent = Cm(0.4)
            paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        else:
            _delete_paragraph(paragraph)
    anchor = paragraphs[min(len(items), len(paragraphs)) - 1]
    for item in items[len(paragraphs):]:
        anchor = _insert_paragraph_after(anchor)
        _set_paragraph_text(anchor, f"\u2022 {item}", size=BODY_FONT_SIZE)
        anchor.paragraph_format.left_indent = Cm(0.4)
        anchor.paragraph_format.first_line_indent = Cm(-0.2)


def _populate_career_summary_table(document: Document, rows: List[Dict[str, str]]) -> None:
    table = _find_table_with_marker(document, "{{CAREER_SUMMARY_COMPANY}}") or _find_table_with_marker(document, "{{COMPANY}}")
    if not table:
        return
    if not rows:
        for para in list(document.paragraphs):
            if para.text.strip() == "Career Summary":
                _delete_paragraph(para)
                break
        _delete_table(table)
        return
    template_row = table.rows[1]
    table._tbl.remove(template_row._tr)
    for item in rows:
        row = table.add_row().cells
        values = [item.get("company", ""), item.get("position", ""), item.get("start_date", ""), item.get("end_date", "")]
        for idx, value in enumerate(values):
            _set_paragraph_text(row[idx].paragraphs[0], value, size=10)
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _apply_cell_borders(row[idx], color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
    _style_table_cells(table)
    _align_table_paragraphs(table, WD_ALIGN_PARAGRAPH.CENTER)


def _populate_qualifications_table(document: Document, qualifications: List[Dict[str, str]]) -> None:
    table = _find_table_with_marker(document, "{{QUALIFICATION_DEGREE}}") or _find_table_with_marker(document, "{{QUALIFICATION}}")
    if not table:
        return
    if not qualifications:
        # remove heading above table
        paras = document.paragraphs
        for i, para in enumerate(paras):
            if para.text.strip() == "Qualification":
                _delete_paragraph(para)
                break
        _delete_table(table)
        return
    template_row = table.rows[1]
    table._tbl.remove(template_row._tr)
    for row_data in qualifications:
        row = table.add_row().cells
        for idx, key in enumerate(["qualification", "institution", "end_date"]):
            value = row_data.get('dates') if key == 'end_date' and row_data.get('dates') else row_data.get(key, "")
            _set_paragraph_text(row[idx].paragraphs[0], value, size=10)
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(row[idx], WHITE_FILL)
            _apply_cell_borders(row[idx], color=SOLID_BLACK_BORDER, size=SOLID_BLACK_BORDER_SIZE)
    _style_qualifications_table(table)


def _add_history_block_after(paragraph, entry: Dict[str, Any]):
    title_para = paragraph
    title = f"{entry.get('position', '')}"
    _set_paragraph_text(title_para, title, bold=True, size=10.5, color=ACCENT_DARK)
    title_para.paragraph_format.space_after = Pt(1)

    meta = _insert_paragraph_after(title_para)
    meta_bits = [entry.get("company", ""), entry.get("dates") or " – ".join([x for x in [entry.get("start_date"), entry.get("end_date")] if x])]
    _set_paragraph_text(meta, " · ".join([_clean(x) for x in meta_bits if _clean(x)]), size=9, color=MUTED)
    meta.paragraph_format.space_after = Pt(2)

    anchor = meta
    if _clean(entry.get("summary")):
        anchor = _insert_paragraph_after(anchor)
        _set_paragraph_text(anchor, _clean(entry.get("summary")), size=9.5, italic=True)
    for bullet_text in entry.get("responsibilities", [])[:8]:
        anchor = _insert_paragraph_after(anchor)
        _set_paragraph_text(anchor, f"• {bullet_text}", size=9.5)
        anchor.paragraph_format.left_indent = Cm(0.45)
        anchor.paragraph_format.first_line_indent = Cm(-0.2)
        anchor.paragraph_format.space_after = Pt(1)
    spacer = _insert_paragraph_after(anchor)
    _set_paragraph_text(spacer, "", size=6)
    return spacer


def _populate_career_history(document: Document, entries: List[Dict[str, Any]]) -> None:
    marker_para = _find_paragraph(document, "{{CAREER_HISTORY_ITEM}}")
    template_table = _find_table_with_marker(document, "{{CAREER_ROLE_COMPANY}}")
    if not template_table:
        return
    if not entries:
        if marker_para:
            _remove_heading_for_marker(document, "{{CAREER_HISTORY_ITEM}}")
        else:
            for paragraph in list(document.paragraphs):
                if paragraph.text.strip() == "Career History":
                    _delete_paragraph(paragraph)
                    break
        _delete_table(template_table)
        return
    if marker_para:
        _delete_paragraph(marker_para)
    pristine_template = deepcopy(template_table._tbl)
    current_table = template_table
    total = len(entries)
    for index, entry in enumerate(entries):
        _fill_career_history_table(current_table, entry)
        if index < total - 1:
            spacer = _insert_paragraph_after_table(current_table)
            spacer.paragraph_format.space_after = Pt(2.5)
            current_table = _insert_table_after_paragraph(spacer, pristine_template)


def build_profile_docx_from_schema(document_path: Path, payload: CandidateProfileSchema | Dict[str, Any]) -> FinalProfilePayload:
    if isinstance(payload, dict):
        payload = CandidateProfileSchema.model_validate(payload)
    final_payload = FinalProfilePayload(
        identity={
            "full_name": payload.identity.full_name,
            "headline": payload.identity.headline,
            "availability": payload.identity.availability or "",
            "region": payload.identity.region or "",
        },
        summary=payload.career_summary,
        skills=_compact_skill_rows([{"category": "", "items": item} for item in payload.skills]),
        qualifications=[{"qualification": q.degree, "institution": q.institution, "end_date": normalize_recruiter_date_text(q.year or "")} for q in payload.qualifications],
        certifications=[
            {
                "name": _normalize_bullet_item(c.cert_name),
                "provider": _normalize_bullet_item(c.provider or ""),
                "year": _normalize_bullet_item(c.year or ""),
            }
            for c in payload.certifications
        ],
        career_history=[
            {
                "position": entry.job_title,
                "company": entry.company,
                "dates": entry.dates,
                "start_date": _parse_date_parts(entry.dates)[0],
                "end_date": _parse_date_parts(entry.dates)[1],
                "responsibilities": entry.bullets,
                "client_engagements": entry.client_engagements,
                "projects": [{"name": project.name, "details": project.details or ""} for project in entry.projects],
                "summary": "",
            }
            for entry in payload.career_history
        ],
    )
    final_payload.career_summary = _career_summary_from_history(final_payload.career_history)
    template_path = _ensure_master_template()
    document = Document(str(template_path))
    _populate_simple_placeholder(document, "{{FULL_NAME}}", final_payload.identity["full_name"])
    _populate_simple_placeholder(document, "{{HEADLINE}}", final_payload.identity["headline"])
    _populate_simple_placeholder(document, "{{PROFESSIONAL_TITLE}}", final_payload.identity["headline"])
    _populate_simple_placeholder(document, "{{CANDIDATE_SUMMARY}}", final_payload.summary)
    _populate_meta_lines(document, final_payload.identity)
    _populate_list_block(document, "{{SKILLS_BLOCK}}", _skill_display_lines(final_payload.skills), bullet=True)
    _populate_qualifications_table(document, final_payload.qualifications)
    _populate_certifications_block(document, [" | ".join([part for part in [c.get('name'), c.get('provider'), c.get('year')] if part]) for c in final_payload.certifications])
    _populate_career_summary_table(document, final_payload.career_summary)
    _populate_career_history(document, final_payload.career_history)
    _set_section_headings_black(document)
    document.save(str(document_path))
    return final_payload


def build_profile_docx(document_path: Path, state: Dict[str, str], profile: Optional[Dict[str, Any]] = None) -> FinalProfilePayload:
    payload = build_final_profile_payload(state, profile)
    validated = validate_export_payload(payload)
    return build_profile_docx_from_schema(document_path, validated)


# ---------------------------------------------------------------------------
# Final targeted overrides for date sorting and user-friendly placeholders
# ---------------------------------------------------------------------------

def _date_sort_value(text: str, *, is_end: bool = False) -> tuple[int, int]:
    value = normalize_recruiter_date_text(_clean(text))
    if not value:
        return (0, 0)
    if re.search(r'\b(?:present|current|now)\b', value, re.I):
        return (9999, 12)
    year_match = re.search(r'((?:19|20)\d{2})', value)
    year = int(year_match.group(1)) if year_match else 0
    month = 12 if is_end else 1
    lowered = value.lower()
    for needle, val in {'jan':1,'january':1,'feb':2,'february':2,'mar':3,'march':3,'apr':4,'april':4,'may':5,'jun':6,'june':6,'jul':7,'july':7,'aug':8,'august':8,'sep':9,'sept':9,'september':9,'oct':10,'october':10,'nov':11,'november':11,'dec':12,'december':12}.items():
        if re.search(rf'\b{needle}\b', lowered):
            month = val
            break
    return (year, month)


def _entry_sort_key(entry: Dict[str, Any]) -> tuple[tuple[int, int], tuple[int, int], str, str]:
    return (
        _date_sort_value(entry.get('end_date') or entry.get('dates') or '', is_end=True),
        _date_sort_value(entry.get('start_date') or entry.get('dates') or ''),
        _clean(entry.get('position')),
        _clean(entry.get('company')),
    )


def _qualification_sort_key(entry: Dict[str, str]) -> tuple[tuple[int, int], str]:
    return (_date_sort_value(entry.get('end_date') or '', is_end=True), _clean(entry.get('qualification')))


def _qualifications_from_state(state: Dict[str, str], profile: Optional[Dict[str, Any]]) -> List[Dict[str, str]]:
    rows = []
    for parts in _split_structured_lines(state.get('education', '')):
        qualification = parts[0] if len(parts) > 0 else ''
        institution = parts[1] if len(parts) > 1 else ''
        start_date = format_recruiter_date(parts[2]) if len(parts) > 3 else ''
        end_date = format_recruiter_date(parts[3]) if len(parts) > 3 else (format_recruiter_date(parts[2]) if len(parts) > 2 else '')
        date_text = ' – '.join([item for item in [start_date, end_date] if item]) if start_date and end_date and start_date != end_date else (end_date or start_date)
        rows.append({
            'qualification': qualification,
            'institution': institution,
            'end_date': end_date,
            'dates': date_text,
        })
    rows = [row for row in rows if row['qualification'] and row['institution']]
    if not rows:
        rows = []
        for item in (profile or {}).get('education', []):
            start_date = format_recruiter_date(item.get('start_date'))
            end_date = normalize_recruiter_date_text(item.get('end_date'))
            date_text = ' – '.join([value for value in [start_date, end_date] if value]) if start_date and end_date and start_date != end_date else (end_date or start_date)
            row = {
                'qualification': _clean(item.get('qualification')),
                'institution': _clean(item.get('institution')),
                'end_date': end_date,
                'dates': date_text,
            }
            if row['qualification'] and row['institution']:
                rows.append(row)
    return sorted(rows, key=_qualification_sort_key, reverse=True)


def _career_history_from_state(state: Dict[str, str], profile: Optional[Dict[str, Any]]) -> List[Dict[str, Any]]:
    parsed = parse_experience_section(state.get('career_history', '') or '')
    profile_entries = list((profile or {}).get('experience', []))
    if not parsed:
        parsed = list(profile_entries)
    elif profile_entries:
        seen = {
            (
                _clean(entry.get('company')).casefold(),
                _clean(entry.get('position')).casefold(),
                _clean(entry.get('start_date')).casefold(),
                _clean(entry.get('end_date')).casefold(),
            )
            for entry in parsed
        }
        for entry in profile_entries:
            key = (
                _clean(entry.get('company')).casefold(),
                _clean(entry.get('position')).casefold(),
                _clean(entry.get('start_date')).casefold(),
                _clean(entry.get('end_date')).casefold(),
            )
            if key not in seen:
                parsed.append(entry)
                seen.add(key)
    profile_index = {
        (
            _clean(entry.get('company')).casefold(),
            _clean(entry.get('position')).casefold(),
            format_recruiter_date(entry.get('start_date')).casefold(),
            format_recruiter_date(entry.get('end_date')).casefold(),
        ): entry
        for entry in profile_entries
    }
    entries: List[Dict[str, Any]] = []
    for entry in parsed:
        company = _clean(entry.get('company'))
        position = _clean(entry.get('position'))
        start_date = format_recruiter_date(entry.get('start_date'))
        end_date = format_recruiter_date(entry.get('end_date'))
        if not position or not company:
            continue
        source_entry = profile_index.get((company.casefold(), position.casefold(), start_date.casefold(), end_date.casefold()), entry)
        clients = list(source_entry.get('clients', []) or [])
        client_engagements = _dedupe_preserve_order(_client_engagement_line(client) for client in clients if _client_engagement_line(client))
        projects = _project_rows_from_clients(clients)
        responsibilities = _dedupe_preserve_order(entry.get('responsibilities', []))[:8]
        if client_engagements:
            responsibilities = [item for item in responsibilities if not _looks_like_client_engagement(item)]
        summary_line = _clean(entry.get('summary'))
        if summary_line and summary_line not in responsibilities:
            responsibilities = [summary_line] + responsibilities
        entries.append(
            {
                'position': position,
                'company': company,
                'start_date': start_date,
                'end_date': end_date,
                'dates': ' – '.join([x for x in [start_date, end_date] if x]),
                'responsibilities': responsibilities[:8],
                'client_engagements': client_engagements[:8],
                'projects': projects[:4],
                'summary': summary_line,
            }
        )
    return sorted(entries, key=_entry_sort_key, reverse=True)


def _career_summary_from_history(entries: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for entry in sorted(entries, key=_entry_sort_key, reverse=True):
        company = _clean(entry.get('company'))
        position = _clean(entry.get('position'))
        start_date = format_recruiter_date(entry.get('start_date')) or _parse_date_parts(entry.get('dates', ''))[0]
        end_date = format_recruiter_date(entry.get('end_date')) or _parse_date_parts(entry.get('dates', ''))[1]
        if company and position:
            rows.append({'company': company, 'position': position, 'start_date': start_date, 'end_date': end_date})
    return rows
