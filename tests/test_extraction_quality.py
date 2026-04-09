from io import BytesIO
from pathlib import Path

from app.pdf_compat import fitz
from docx import Document
from fastapi.testclient import TestClient

from app.main import create_app
from app.normalizers import profile_from_sections, profile_to_template_state, validate_profile_readiness
from app.parsers import extract_identity, parse_sections
from app.utils_text import extract_text


def test_docx_table_extraction_normalizes_qualifications_and_certifications(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Thandokuhle Mntambo")
    doc.add_paragraph("QA Analyst")
    doc.add_paragraph("Qualifications")
    q_table = doc.add_table(rows=3, cols=3)
    q_table.rows[0].cells[0].text = "Qualification"
    q_table.rows[0].cells[1].text = "Institution"
    q_table.rows[0].cells[2].text = "Year"
    q_table.rows[1].cells[0].text = "BSc Information Technology"
    q_table.rows[1].cells[1].text = "North-West University"
    q_table.rows[1].cells[2].text = "2024"
    q_table.rows[2].cells[0].text = "Advanced Diploma Quality Assurance"
    q_table.rows[2].cells[1].text = "Test Institute"
    q_table.rows[2].cells[2].text = "2025"
    doc.add_paragraph("Certifications")
    c_table = doc.add_table(rows=2, cols=3)
    c_table.rows[0].cells[0].text = "Certification"
    c_table.rows[0].cells[1].text = "Provider"
    c_table.rows[0].cells[2].text = "Year"
    c_table.rows[1].cells[0].text = "ISTQB Foundation"
    c_table.rows[1].cells[1].text = "ISTQB"
    c_table.rows[1].cells[2].text = "2025"
    path = tmp_path / "table_cv.docx"
    doc.save(path)

    raw_text = extract_text(path)
    sections = parse_sections(raw_text)
    profile = profile_from_sections(raw_text, sections, path)

    assert any(item["qualification"] == "BSc Information Technology" for item in profile["education"])
    assert any("ISTQB" in cert for cert in profile["certifications"])


def test_pdf_table_like_extraction_normalizes_identity_and_experience(tmp_path: Path):
    path = tmp_path / "candidate.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    y = 40
    for line in [
        "Thandokuhle Mntambo",
        "Software Tester",
        "Email: thando@example.com",
        "Career History",
        "Company            Role                              Start Date      End Date",
        "OpenText           QA Intern                         Jan 2025        Present",
        "Validated web and API workflows across Agile delivery teams.",
    ]:
        page.insert_text((40, y), line, fontsize=11)
        y += 18
    pdf.save(path)
    pdf.close()

    raw_text = extract_text(path)
    sections = parse_sections(raw_text)
    profile = profile_from_sections(raw_text, sections, path)

    assert profile["identity"]["full_name"] == "Thandokuhle Mntambo"
    assert profile["identity"]["headline"] == "Software Tester"
    assert any(entry.get("company") == "OpenText" for entry in profile["experience"])
    assert any(entry.get("position") == "QA Intern" for entry in profile["experience"])


def test_references_do_not_contaminate_identity(tmp_path: Path):
    raw_text = """
    Email: candidate@example.com
    Phone: 012 345 6789

    References
    Dr Jane Smith
    Senior Lecturer
    jane@example.com
    """
    sections = parse_sections(raw_text)
    identity = extract_identity(raw_text, sections, tmp_path / "Thandokuhle_Mntambo_CV.docx")
    assert identity["full_name"] == "Thandokuhle Mntambo"
    assert identity["full_name"] != "Dr Jane Smith"


def test_weak_career_summary_is_rejected():
    state = {
        "full_name": "Thandokuhle Mntambo",
        "headline": "QA Analyst",
        "summary": "I am a hardworking team player looking for growth.",
        "skills": "Selenium\nAPI Testing",
        "education": "BSc IT | NWU | 2024",
        "career_history": "QA Intern - OpenText Jan 2025 - Present\nExecuted regression tests.",
    }
    issues = validate_profile_readiness(state)
    assert any("Career Summary" in issue for issue in issues)


def test_review_is_reset_after_template_change():
    app = create_app()
    client = TestClient(app)
    doc = Document()
    for line in [
        "George Thabiso Mpopo",
        "Senior Software Engineer",
        "Summary",
        "Experienced software engineer delivering enterprise systems across regulated environments with strong implementation and delivery capability.",
        "Skills",
        "C#, .NET, Azure",
        "Qualifications",
        "BSc Computer Science | University of Zululand | 2007",
        "Career History",
        "Gijima Technologies | Senior Software Engineer | 2024 | Present",
        "Delivered enterprise-grade platforms for public-sector clients.",
    ]:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    upload = client.post("/api/upload", files={"file": ("candidate.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    payload = upload.json()

    review = client.post(f"/api/document/{payload['document_id']}/review-complete", json={"template_state": payload["template_state"]})
    assert review.status_code == 200

    update = client.post(f"/api/document/{payload['document_id']}/template", json={"summary": "I am looking for a job."})
    assert update.status_code == 200
    assert update.json()["workflow_state"]["can_download"] is False

    download = client.post(f"/api/document/{payload['document_id']}/download", json={})
    assert download.status_code == 400


def test_lindelwe_summary_name_and_paragraph_skills_are_preserved_from_source():
    path = Path("/mnt/data/Lindelwe Myeza Resume 2025.pdf")
    source_path = path
    if not path.exists():
        path = Path(__file__).resolve().parents[1] / "uploads" / "fcd22d91-3baf-42a5-b10f-b06199029a11.pdf"
        source_path = Path("LINDELWE MYEZA Resume 2025.pdf")
    raw_text = extract_text(path)
    sections = parse_sections(raw_text)
    profile = profile_from_sections(raw_text, sections, source_path)
    state = profile_to_template_state(profile)

    using_original_source = source_path == Path("/mnt/data/Lindelwe Myeza Resume 2025.pdf")
    expected_name = "LINDELWE MYEZA" if using_original_source else "Lindelwe Myeza"
    assert profile["identity"]["full_name"] == expected_name
    assert state["summary"].startswith("Currently my short-term objectives are to gain new skills")
    expected_summary_fragment = (
        "Africa become a world class competitor"
        if using_original_source
        else "contributing effectively as a COBOL software developer"
    )
    assert expected_summary_fragment in state["summary"]
    assert "Critical thinking" in state["skills"]
    assert "Problem-solving" in state["skills"]
    assert state["summary"].count("\n") == 0
    assert state["availability"] == "Availability not provided"
    assert state["region"] == "Region not provided"


def test_date_driven_tables_sort_most_recent_first_and_use_responsibilities_label(tmp_path: Path):
    raw = """
    Thandokuhle Mntambo
    QA Analyst
    Objective
    I am looking to build my testing career through structured delivery work.
    Skills
    API testing, Jira
    Qualifications
    National Senior Certificate | Queens High School | 2018
    BCom Information Systems | University of Johannesburg | 2021
    Career History
    OpenText | QA Analyst | Feb 2021 | Present
    Executed regression testing.
    Acme Corp | QA Intern | Jun 2020 | Jan 2021
    Supported smoke testing.
    """
    profile = profile_from_sections(raw, parse_sections(raw), tmp_path / "candidate.docx")
    state = profile_to_template_state(profile)
    assert state["career_summary"].splitlines()[0].startswith("QA Analyst - OpenText")
    assert state["education"].splitlines()[0].endswith("2021")
    assert "Responsibilities:" in state["career_history"]
