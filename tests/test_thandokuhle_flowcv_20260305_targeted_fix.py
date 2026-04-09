from pathlib import Path

from docx import Document

from app.docx_exporter import build_final_profile_payload, build_profile_docx_from_schema
from app.models import build_detected_blocks
from app.normalizers import profile_from_sections, profile_to_template_state
from app.parsers import build_source_sections, format_recruiter_date, normalize_recruiter_date_text, parse_sections
from app.renderers import build_preview_html
from app.schema import validate_export_payload


FLOWCV_RAW = """
Thandokuhle Mnqobi Mntambo
QA Intern mnqobimntambo@gmail.com

078 568 3003

Bryanston, South Africa

za.linkedin.com/in/thandokuhle-mntambo

Professional Summary
Information Systems Honours graduate with strong skills in SQL, Power BI, data analytics, and business intelligence.
Experienced in transforming raw data into actionable insights through data modelling, visualization, and reporting.
Certified in Microsoft Azure Data Fundamentals and Google Data Analytics, with a strong interest in building data-driven solutions that support business decision-making.
Professional Experience
QA Intern, Cestasoft Solutions
\u2022Designed and executed manual and automated test cases using UFT One to validate application functionality, reliability, and data integrity.
06/2025 \u2013 Present
Bryanston
\u2022Logged, tracked, and prioritised defects using QMetry and Jira while maintaining detailed bug reports and test documentation
\u2022Analysed system requirements and prepared test data to ensure accurate test coverage and validation of business logic.
\u2022Worked within Agile and DevOps workflows, supporting CI/ CD pipelines and collaborative software delivery.
\u2022Applied knowledge of the Software Testing Life Cycle (STLC), including test design, execution, defect management, and automation practices.
Education
BCom Honours - Information Systems and Business Management,
Nelson Mandela University
2023 \u2013 2024
BCom - Information Systems and Business Management, Nelson Mandela University 2020 \u2013 2022
Technical Skills
Data & Analytics
\u2022Power BI
\u2022Data Cleaning & Preparation
\u2022DAX basics
\u2022Data Modelling
\u2022SQL (Microsoft SQL Server)
\u2022Data Visualization & Reporting
Programming
\u2022Python
\u2022C# (ASP.NET fundamentals)
\u2022HTML/ CSS
\u2022Java
Professional Skills
\u2022Analytical Thinking & Problem Solving
\u2022Communication & Team Collaboration
\u2022Time Management & Organization
\u2022Attention to Detail
Data Engineering Foundations
\u2022Data Warehousing Concepts
\u2022ETL Fundamentals & Data Transformation
Productivity Tools
\u2022Microsoft Excel (data analysis & visualization)
\u2022Microsoft Office Suite
Automation Testing
\u2022UFT One
\u2022Selenium WebDriver
Projects
Student Feedback System (NexusEd) - Honours Project
\u2022Conducted stakeholder interviews, surveys and requirements gathering
02/2024 \u2013 10/2024
\u2022Designed system architecture using UML diagrams including use case, activity and ER diagrams
\u2022Developed a web-based student feedback platform using C# (ASP.NET) and HTML/ CSS
\u2022Connected the application to a relational SQL database
\u2022Used SQL queries to extract, clean and analyse system-generated data
\u2022Created visual analytics dashboards using Chart.js
\u2022Produced a formal research thesis documenting system design and insights
Certifications
OpenText Certified Developer -
Unified Functional Testing One
(UFT One)
Google Data Analytics Professional
Certificate
Microsoft Azure Data
Fundamentals (DP-900)
Microsoft SQL Server
Google Business Intelligence
Professional Certificate
Microsoft Power BI Data Analyst
Associate (In Progress)
Courses
Google Data Analysis with Python, Coursera
Microsoft SQL Server, Coursera
""".strip()


def test_flowcv_20260305_normalizes_certifications_skills_and_experience(tmp_path: Path):
    profile = profile_from_sections(
        FLOWCV_RAW,
        parse_sections(FLOWCV_RAW),
        tmp_path / "Thandokuhle_Mnqobi_Mntambo_FlowCV_Resume_2026-03-05 (3).pdf",
    )
    state = profile_to_template_state(profile)

    assert profile["identity"]["full_name"] == "Thandokuhle Mnqobi Mntambo"
    assert profile["identity"]["headline"] == "QA Intern"
    assert profile["certifications"] == [
        "OpenText Certified Developer - Unified Functional Testing One (UFT One)",
        "Google Data Analytics Professional Certificate",
        "Microsoft Azure Data Fundamentals (DP-900)",
        "Microsoft SQL Server",
        "Google Business Intelligence Professional Certificate",
        "Microsoft Power BI Data Analyst Associate (In Progress)",
    ]
    assert state["skills"].splitlines() == [
        "Data & Analytics",
        "Power BI",
        "Data Cleaning & Preparation",
        "DAX basics",
        "Data Modelling",
        "SQL (Microsoft SQL Server)",
        "Data Visualization & Reporting",
        "Programming",
        "Python",
        "C# (ASP.NET fundamentals)",
        "HTML/ CSS",
        "Java",
        "Professional Skills",
        "Analytical Thinking & Problem Solving",
        "Communication & Team Collaboration",
        "Time Management & Organization",
        "Attention to Detail",
        "Data Engineering Foundations",
        "Data Warehousing Concepts",
        "ETL Fundamentals & Data Transformation",
        "Productivity Tools",
        "Microsoft Excel (data analysis & visualization)",
        "Microsoft Office Suite",
        "Automation Testing",
        "UFT One",
        "Selenium WebDriver",
    ]
    assert "Programming Languages:" not in state["skills"]
    assert profile["experience"] == [
        {
            "company": "Cestasoft Solutions",
            "position": "QA Intern",
            "start_date": "Jun 2025",
            "end_date": "Present",
            "responsibilities": [
                "Designed and executed manual and automated test cases using UFT One to validate application functionality, reliability, and data integrity.",
                "Logged, tracked, and prioritised defects using QMetry and Jira while maintaining detailed bug reports and test documentation",
                "Analysed system requirements and prepared test data to ensure accurate test coverage and validation of business logic.",
                "Worked within Agile and DevOps workflows, supporting CI/ CD pipelines and collaborative software delivery.",
                "Applied knowledge of the Software Testing Life Cycle (STLC), including test design, execution, defect management, and automation practices.",
            ],
            "clients": [],
            "technologies": [],
            "summary": None,
        }
    ]
    assert state["career_summary"] == "QA Intern - Cestasoft Solutions (Jun 2025 – Present)"
    assert "QA Intern - Cestasoft Solutions (Jun 2025 – Present)" in state["career_history"]
    assert "06/2025" not in state["career_summary"]
    assert "06/2025" not in state["career_history"]
    assert "QMetry and Jira" in state["career_history"]


def test_flowcv_20260305_identity_header_is_not_reused_as_additional_information():
    sections = parse_sections(FLOWCV_RAW)

    assert all(section["mapped_field"] != "additional_sections" for section in build_source_sections(sections))
    assert all(block["mapped_field"] != "additional_sections" for block in build_detected_blocks(sections))


def test_flowcv_20260305_preview_and_docx_use_abbreviated_month_dates(tmp_path: Path):
    profile = profile_from_sections(
        FLOWCV_RAW,
        parse_sections(FLOWCV_RAW),
        tmp_path / "Thandokuhle_Mnqobi_Mntambo_CV.pdf",
    )
    state = profile_to_template_state(profile)

    preview = build_preview_html(state, profile)
    assert "Jun 2025" in preview
    assert "Jun 2025 – Present" in preview
    assert "06/2025" not in preview

    payload = validate_export_payload(build_final_profile_payload(state, profile))
    output = tmp_path / "thandokuhle_flowcv_output.docx"
    build_profile_docx_from_schema(output, payload)

    doc = Document(output)
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()) + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()
    )
    assert "Jun 2025" in full_text
    assert "Present" in full_text
    assert "06/2025" not in full_text


def test_recruiter_date_formatters_normalize_month_year_and_ranges():
    assert format_recruiter_date("06/2025") == "Jun 2025"
    assert format_recruiter_date("10/2024") == "Oct 2024"
    assert format_recruiter_date("February 2024") == "Feb 2024"
    assert format_recruiter_date("2023") == "2023"
    assert format_recruiter_date("Present") == "Present"
    assert normalize_recruiter_date_text("02/2024 - 10/2024") == "Feb 2024 – Oct 2024"
    assert normalize_recruiter_date_text("2023 - 2024") == "2023 – 2024"
