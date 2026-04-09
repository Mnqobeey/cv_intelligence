from pathlib import Path

from docx import Document

from app.docx_exporter import build_final_profile_payload
from app.normalizers import profile_from_sections, profile_to_template_state, validate_profile_readiness
from app.docx_exporter import build_profile_docx_from_schema
from app.parsers import parse_sections
from app.renderers import build_preview_html
from app.schema import validate_export_payload


CONSULTING_CV_6_RAW = """
Tshepo Ndlovu
DevOps Engineer | Johannesburg, South Africa
Email: consultant@example.com | Phone: +27 82 000 0000
Professional Summary
Consulting professional delivering enterprise technology solutions through consulting partners for banking, telecom, fintech, insurance, and government clients. Experienced in system integration, cloud deployment, and enterprise software delivery.
Core Skills
Python, Linux, Azure, Terraform, Kubernetes, Docker, Kafka, Selenium
Professional Experience
DevOps Engineer – CestaSoft Solutions (Consulting Partner)
2020 – Present
Client Projects Delivered:
Client: Nedbank
Delivered enterprise platform enhancements including API integrations, automation frameworks, and cloud infrastructure improvements supporting high-volume enterprise environments.
Client: Discovery Insurance
Delivered enterprise platform enhancements including API integrations, automation frameworks, and cloud infrastructure improvements supporting high-volume enterprise environments.
Client: FNB
Delivered enterprise platform enhancements including API integrations, automation frameworks, and cloud infrastructure improvements supporting high-volume enterprise environments.
Client: Absa Bank
Delivered enterprise platform enhancements including API integrations, automation frameworks, and cloud infrastructure improvements supporting high-volume enterprise environments.
Client: National Treasury
Delivered enterprise platform enhancements including API integrations, automation frameworks, and cloud infrastructure improvements supporting high-volume enterprise environments.
Previous Experience
Software Developer – DigitalTech Labs
2017 – 2020
• Developed backend services and enterprise integrations.
• Implemented CI/CD pipelines and automated testing frameworks.
Education
BSc Information Technology – University of Johannesburg
Certifications
Microsoft Azure Fundamentals
AWS Cloud Practitioner
""".strip()


def _build_profile(tmp_path: Path):
    sections = parse_sections(CONSULTING_CV_6_RAW)
    profile = profile_from_sections(CONSULTING_CV_6_RAW, sections, tmp_path / "consulting_cv_6.docx")
    state = profile_to_template_state(profile)
    return profile, state


def test_consulting_cv6_extracts_identity_skills_and_optional_portfolio_cleanly(tmp_path: Path):
    profile, state = _build_profile(tmp_path)
    issues = validate_profile_readiness(state)

    assert profile["identity"]["full_name"] == "Tshepo Ndlovu"
    assert profile["identity"]["headline"] == "DevOps Engineer"
    assert profile["identity"]["location"] == "Johannesburg, South Africa"
    assert profile["identity"]["email"] == "consultant@example.com"
    assert profile["identity"]["phone"] == "+27 82 000 0000"

    assert state["portfolio"] == "Portfolio not provided"
    assert state["skills"].splitlines() == [
        "Python",
        "Linux",
        "Azure",
        "Terraform",
        "Kubernetes",
        "Docker",
        "Kafka",
        "Selenium",
    ]
    assert state["education"] == "BSc Information Technology | University of Johannesburg"
    assert state["certifications"].splitlines() == [
        "Microsoft Azure Fundamentals",
        "AWS Cloud Practitioner",
    ]
    assert not any("Qualifications contain malformed" in issue for issue in issues)
    assert not any("Career History" in issue for issue in issues)
    assert not any("Portfolio / Website" in issue for issue in issues)


def test_consulting_cv6_keeps_single_parent_consulting_role_and_previous_role_separate(tmp_path: Path):
    profile, state = _build_profile(tmp_path)

    assert len(profile["experience"]) == 2

    current_role = profile["experience"][0]
    previous_role = profile["experience"][1]

    assert current_role["position"] == "DevOps Engineer"
    assert current_role["company"] == "CestaSoft Solutions (Consulting Partner)"
    assert current_role["start_date"] == "2020"
    assert current_role["end_date"] == "Present"
    assert [client["client_name"] for client in current_role["clients"]] == [
        "Nedbank",
        "Discovery Insurance",
        "FNB",
        "Absa Bank",
        "National Treasury",
    ]
    assert all(client["responsibilities"] for client in current_role["clients"])
    assert all(client["client_name"] not in {entry["company"] for entry in profile["experience"]} for client in current_role["clients"])

    assert previous_role["position"] == "Software Developer"
    assert previous_role["company"] == "DigitalTech Labs"
    assert previous_role["start_date"] == "2017"
    assert previous_role["end_date"] == "2020"

    payload = build_final_profile_payload(state, profile)
    assert payload.career_summary == [
        {
            "company": "CestaSoft Solutions (Consulting Partner)",
            "position": "DevOps Engineer",
            "start_date": "2020",
            "end_date": "Present",
        },
        {
            "company": "DigitalTech Labs",
            "position": "Software Developer",
            "start_date": "2017",
            "end_date": "2020",
        },
    ]
    assert [entry["company"] for entry in payload.career_history] == [
        "CestaSoft Solutions (Consulting Partner)",
        "DigitalTech Labs",
    ]
    assert payload.career_history[0]["responsibilities"] == []
    assert any("Nedbank:" in item for item in payload.career_history[0]["client_engagements"])
    assert any("National Treasury:" in item for item in payload.career_history[0]["client_engagements"])
    assert payload.career_history[0]["projects"] == []
    assert state["career_history"].count("DevOps Engineer - CestaSoft Solutions (Consulting Partner) (2020 – Present)") == 1
    assert state["career_history"].count("Software Developer - DigitalTech Labs (2017 – 2020)") == 1


def test_consulting_cv6_docx_uses_client_engagement_heading_and_vertical_skills(tmp_path: Path):
    profile, state = _build_profile(tmp_path)
    validated = validate_export_payload(build_final_profile_payload(state, profile))
    output = tmp_path / "consulting_cv_6_output.docx"

    build_profile_docx_from_schema(output, validated)

    doc = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    full_text = paragraph_text + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()
    )
    detail_tables = [table for table in doc.tables if any("Company Name" in cell.text for row in table.rows for cell in row.cells)]

    assert "Core Skills:" not in full_text
    assert "Python\nLinux\nAzure\nTerraform\nKubernetes\nDocker\nKafka\nSelenium" in paragraph_text
    assert "Python, Linux, Azure, Terraform, Kubernetes, Docker, Kafka, Selenium" not in full_text
    assert len(detail_tables) == 2

    consulting_detail = detail_tables[0].rows[3].cells[0].text
    standard_detail = detail_tables[1].rows[3].cells[0].text

    assert "Client Engagements:" in consulting_detail
    assert "Responsibilities:" not in consulting_detail
    assert "Nedbank:" in consulting_detail
    assert "Discovery Insurance:" in consulting_detail
    assert "FNB:" in consulting_detail
    assert "Absa Bank:" in consulting_detail
    assert "National Treasury:" in consulting_detail
    assert "Responsibilities:" in standard_detail
    assert "Client Engagements:" not in standard_detail


def test_consulting_cv6_preview_uses_client_engagement_heading_for_consulting_role(tmp_path: Path):
    profile, state = _build_profile(tmp_path)

    preview = build_preview_html(state, profile)

    assert "Client Engagements" in preview
    assert "Responsibilities" in preview
    assert preview.count("Responsibilities") == 1
    assert "Nedbank: Delivered enterprise platform enhancements" in preview
    assert "Discovery Insurance: Delivered enterprise platform enhancements" in preview
    assert "FNB: Delivered enterprise platform enhancements" in preview
    assert "Absa Bank: Delivered enterprise platform enhancements" in preview
    assert "National Treasury: Delivered enterprise platform enhancements" in preview
