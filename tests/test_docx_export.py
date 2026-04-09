from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.docx_exporter import MASTER_TEMPLATE_PATH, build_final_profile_payload, build_profile_docx, build_profile_docx_from_schema
from app.normalizers import profile_to_template_state
from app.schema import CandidateProfileSchema


def sample_state():
    return {
        "full_name": "George Thabiso Mpopo",
        "headline": "Senior / Lead Software Engineer",
        "availability": "Immediate / Negotiable",
        "region": "South Africa",
        "summary": (
            "Senior / Lead Software Engineer with 15+ years of experience delivering secure, scalable, "
            "enterprise-grade systems across regulated industries."
        ),
        "skills": "Languages & Frameworks: C#, .NET, ASP.NET Core\nFrontend: Angular, TypeScript, JavaScript\nCloud / DevOps: Azure DevOps, Docker, Kubernetes",
        "education": "BSc Degree in Computer Science | University of Zululand | 2007\nPostgraduate Diploma in Business Management | MANCOSA | In Progress",
        "certifications": "TOGAF 9.2 Certification (2018)\nIBM Spectrum Protect Version 1.8 (2018)",
        "career_history": (
            "Senior Full-Stack Software Developer - Gijima Technologies Apr 2024 - Present\n"
            "Lead full-stack development on national-scale government systems supporting high-volume citizen services.\n"
            "Design Angular frontends and secure ASP.NET Core REST APIs aligned to enterprise standards.\n\n"
            "Senior Full-Stack Software Developer - Investec Bank Limited Feb 2022 - Apr 2024\n"
            "Delivered mission-critical banking platforms supporting high-value financial transactions.\n"
            "Developed Angular applications for internal teams and external banking users."
        ),
    }


def test_final_payload_repairs_combined_role_and_company():
    payload = build_final_profile_payload(sample_state(), profile=None)
    assert payload.identity["full_name"] == "George Thabiso Mpopo"
    assert payload.career_history[0]["company"] == "Gijima Technologies"
    assert payload.career_history[0]["position"] == "Senior Full-Stack Software Developer"
    assert payload.career_history[1]["company"] == "Investec Bank Limited"


def test_docx_export_contains_expected_sections_and_uses_master_template(tmp_path: Path):
    output = tmp_path / "profile.docx"
    build_profile_docx(output, sample_state(), profile=None)
    assert output.exists()
    assert MASTER_TEMPLATE_PATH.exists()

    doc = Document(output)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    full_text = text + "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "George Thabiso Mpopo" in full_text
    assert "Candidate Summary" in full_text
    assert "Career Summary" in full_text
    assert "Qualification" in full_text
    assert "Career History" in full_text
    assert "{{" not in full_text
    assert "Gijima Technologies" in full_text


def test_template_export_omits_optional_meta_and_certifications_cleanly(tmp_path: Path):
    payload = CandidateProfileSchema.model_validate(
        {
            "identity": {"full_name": "Jane Doe", "headline": "QA Analyst", "availability": "", "region": ""},
            "career_summary": "QA Analyst with experience supporting manual and structured test delivery across enterprise applications.",
            "skills": ["Testing", "Regression Testing", "JIRA"],
            "qualifications": [{"degree": "BSc IT", "institution": "ABC University", "year": "2024"}],
            "certifications": [],
            "career_history": [{"job_title": "QA Analyst", "company": "OpenText", "dates": "2024 – Present", "bullets": ["Executed regression tests."]}],
        }
    )
    output = tmp_path / "profile_no_optional.docx"
    build_profile_docx_from_schema(output, payload)
    doc = Document(output)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip()) + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "Availability:" not in full_text
    assert "Region:" not in full_text
    assert "Certifications" not in full_text
    assert "Career Summary" in full_text
    assert "Jane Doe" in full_text


def test_career_history_repeats_render_as_stable_role_tables(tmp_path: Path):
    output = tmp_path / "career_tables.docx"
    build_profile_docx(output, sample_state(), profile=None)
    doc = Document(output)
    career_tables = [table for table in doc.tables if any("Company Name" in cell.text for row in table.rows for cell in row.cells)]
    assert len(career_tables) == 2
    companies = [table.rows[0].cells[1].text for table in career_tables]
    assert companies == ["Gijima Technologies", "Investec Bank Limited"]
    second_resp = career_tables[1].rows[3].cells[0].text
    assert "Delivered mission-critical banking platforms" in second_resp
    assert "{{" not in second_resp


def test_template_export_uses_master_template_role_markers_without_leakage(tmp_path: Path):
    payload = CandidateProfileSchema.model_validate(
        {
            "identity": {"full_name": "Jane Doe", "headline": "QA Analyst", "availability": "", "region": ""},
            "career_summary": "QA Analyst with experience supporting manual and structured test delivery across enterprise applications.",
            "skills": ["Testing", "Regression Testing", "JIRA", "SQL", "API Testing", "UAT", "Defect Tracking"],
            "qualifications": [{"degree": "BSc IT", "institution": "ABC University", "year": "2024"}],
            "certifications": [],
            "career_history": [
                {"job_title": "QA Analyst", "company": "OpenText", "dates": "2024 – Present", "bullets": ["Executed regression tests.", "Logged and tracked defects."]},
                {"job_title": "QA Intern", "company": "Acme", "dates": "2023 – 2024", "bullets": ["Supported smoke testing."]},
            ],
        }
    )
    output = tmp_path / "template_population.docx"
    build_profile_docx_from_schema(output, payload)
    doc = Document(output)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip()) + "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "{{CAREER_ROLE_COMPANY}}" not in full_text
    assert "{{CLIENT_ENGAGEMENTS_BLOCK}}" not in full_text
    assert "{{PROJECT_NAME}}" not in full_text
    assert "Certifications" not in full_text
    assert "Career Summary" in full_text
    assert full_text.index("Career History") < full_text.index("OpenText")
    assert full_text.index("OpenText") < full_text.index("Acme")
    assert "George Thabiso Mpopo" not in full_text
    assert "Gijima Technologies" not in full_text


def test_role_detail_renders_optional_client_engagements_and_projects_cleanly(tmp_path: Path):
    payload = CandidateProfileSchema.model_validate(
        {
            "identity": {"full_name": "Jane Doe", "headline": "Consulting Engineer", "availability": "Immediate", "region": "Johannesburg"},
            "career_summary": "Consulting engineer delivering enterprise platform work across client-facing environments.",
            "skills": ["Python", "Azure", "Kubernetes"],
            "qualifications": [{"degree": "BSc IT", "institution": "ABC University", "year": "2024"}],
            "certifications": [{"cert_name": "Azure Fundamentals", "provider": "Microsoft", "year": "2023"}],
            "career_history": [
                {
                    "job_title": "Consulting Engineer",
                    "company": "CestaSoft Solutions (Consulting Partner)",
                    "dates": "2024 â€“ Present",
                    "bullets": ["Led platform delivery governance.", "Owned release readiness and environment stability."],
                    "client_engagements": ["Nedbank: Delivered API integration support.", "FNB: Supported platform hardening."],
                    "projects": [{"name": "Payments Modernisation", "details": "Cloud migration and release coordination."}],
                }
            ],
        }
    )
    output = tmp_path / "role_detail_mix.docx"
    build_profile_docx_from_schema(output, payload)
    doc = Document(output)
    detail_table = next(table for table in doc.tables if any("Company Name" in cell.text for row in table.rows for cell in row.cells))
    detail_text = detail_table.rows[3].cells[0].text

    assert "Responsibilities:" in detail_text
    assert "Client Engagements:" in detail_text
    assert "Projects (Linked to this role):" in detail_text
    assert "Payments Modernisation" in detail_text
    assert "Cloud migration and release coordination." in detail_text
    assert "{{" not in detail_text

def test_export_follows_strict_george_section_order(tmp_path: Path):
    output = tmp_path / "strict_order.docx"
    build_profile_docx(output, sample_state(), profile=None)
    doc = Document(output)
    combined = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    assert combined.index("George Thabiso Mpopo") < combined.index("Candidate Summary")
    assert combined.index("Candidate Summary") < combined.index("Skills")
    assert combined.index("Skills") < combined.index("Qualification")
    assert combined.index("Qualification") < combined.index("Certifications")
    assert combined.index("Certifications") < combined.index("Career Summary")
    assert combined.index("Career Summary") < combined.index("Career History")


def test_experience_always_renders_career_summary_table_and_detail_section(tmp_path: Path):
    output = tmp_path / "experience_sections.docx"
    build_profile_docx(output, sample_state(), profile=None)
    doc = Document(output)
    summary_tables = [t for t in doc.tables if any(cell.text.strip() == "Company" for row in t.rows for cell in row.cells)]
    detail_tables = [t for t in doc.tables if any("Company Name" in cell.text for row in t.rows for cell in row.cells)]
    assert len(summary_tables) == 1
    assert len(detail_tables) == 2
    summary_text = "\n".join(cell.text for row in summary_tables[0].rows for cell in row.cells)
    assert "Gijima Technologies" in summary_text
    assert "Investec Bank Limited" in summary_text


def test_export_applies_arial_template_and_bullet_list_formatting(tmp_path: Path):
    state = sample_state()
    state["certifications"] = "•\t• TOGAF 9.2 Certification | Open Group | 2018\n•\t• IBM Spectrum Protect Version 1.8 | IBM | 2018"

    output = tmp_path / "formatted_profile.docx"
    build_profile_docx(output, state, profile=None)
    doc = Document(output)

    heading_para = next(
        paragraph for paragraph in doc.paragraphs if paragraph.text.strip() in {
            "George Thabiso Mpopo",
            "George Thabiso Mpopo - Senior / Lead Software Engineer",
        }
    )
    role_para = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Senior / Lead Software Engineer"),
        heading_para,
    )
    candidate_summary_para = next(
        (
            paragraph
            for paragraph in doc.paragraphs
            if paragraph.text.strip().startswith("Senior / Lead Software Engineer with 15+ years of experience")
        ),
        doc.tables[0].cell(0, 0).paragraphs[0],
    )
    availability_para = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Availability: Immediate / Negotiable")
    region_para = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Region: South Africa")
    skills_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Skills")
    skills_bullets = []
    collecting_skills = False
    certifications_lines = []
    collecting_certs = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Skills":
            collecting_skills = True
            collecting_certs = False
            continue
        if text == "Qualification":
            collecting_skills = False
        if text == "Certifications":
            collecting_certs = True
            continue
        if text == "Career Summary":
            collecting_certs = False
        if collecting_skills and text:
            skills_bullets.append(text)
        if collecting_certs and text:
            certifications_lines.append(text)

    assert heading_para.runs[0].font.name == "Arial"
    assert round(heading_para.runs[0].font.size.pt) == 18
    assert heading_para.runs[0].bold
    assert role_para.runs[0].font.name == "Arial"
    assert round(role_para.runs[0].font.size.pt) == 18
    assert role_para.runs[0].bold
    assert skills_heading.runs[0].font.name == "Arial"
    assert round(skills_heading.runs[0].font.size.pt) == 11
    assert availability_para.text.strip() == "Availability: Immediate / Negotiable"
    assert availability_para.runs[0].font.name == "Arial"
    assert round(availability_para.runs[0].font.size.pt) == 11
    assert availability_para.runs[0].bold
    assert region_para.text.strip() == "Region: South Africa"
    assert region_para.runs[0].font.name == "Arial"
    assert round(region_para.runs[0].font.size.pt) == 11
    assert region_para.runs[0].bold
    assert candidate_summary_para.runs[0].font.name == "Arial"
    assert round(candidate_summary_para.runs[0].font.size.pt) == 10
    assert candidate_summary_para.alignment == 0
    assert skills_bullets
    assert all(line.startswith("• ") for line in skills_bullets)
    assert certifications_lines == [
        "• TOGAF 9.2 Certification | Open Group | 2018",
        "• IBM Spectrum Protect Version 1.8 | IBM | 2018",
    ]


def test_schema_certifications_with_prefixed_bullets_do_not_double_render(tmp_path: Path):
    payload = CandidateProfileSchema.model_validate(
        {
            "identity": {"full_name": "Jane Doe", "headline": "BI Analyst", "availability": "", "region": ""},
            "career_summary": "Business intelligence analyst delivering reporting, analytics, and dashboard solutions across enterprise environments.",
            "skills": ["Power BI", "SQL"],
            "qualifications": [{"degree": "BSc IT", "institution": "ABC University", "year": "2024"}],
            "certifications": [
                {"cert_name": "•\t• Microsoft Certified Power BI Data Analyst Associate", "provider": "• Microsoft", "year": ""},
                {"cert_name": "•\t• Google Business Intelligence Professional Certificate", "provider": "• Google", "year": ""},
            ],
            "career_history": [{"job_title": "BI Analyst", "company": "Acme", "dates": "2024 – Present", "bullets": ["Built reporting dashboards."]}],
        }
    )
    output = tmp_path / "schema_certifications.docx"
    build_profile_docx_from_schema(output, payload)
    doc = Document(output)
    cert_lines = []
    collecting = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Certifications":
            collecting = True
            continue
        if text == "Career Summary":
            collecting = False
        if collecting and text:
            cert_lines.append(text)
    assert cert_lines == [
        "• Microsoft Certified Power BI Data Analyst Associate | Microsoft",
        "• Google Business Intelligence Professional Certificate | Google",
    ]
    cert_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() in cert_lines]
    assert cert_paragraphs
    for paragraph in cert_paragraphs:
        p_pr = paragraph._p.pPr
        assert p_pr is None or p_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') is None


def test_export_uses_black_identity_text_and_solid_black_table_borders(tmp_path: Path):
    output = tmp_path / "styled_profile.docx"
    build_profile_docx(output, sample_state(), profile=None)
    doc = Document(output)

    heading_para = next(
        paragraph for paragraph in doc.paragraphs if paragraph.text.strip() in {
            "George Thabiso Mpopo",
            "George Thabiso Mpopo - Senior / Lead Software Engineer",
        }
    )
    role_para = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Senior / Lead Software Engineer"),
        heading_para,
    )
    availability_para = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Availability: Immediate / Negotiable")
    region_para = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Region: South Africa")

    for paragraph in [heading_para, role_para, availability_para, region_para]:
        assert str(paragraph.runs[0].font.color.rgb) == "000000"

    for label in ["Candidate Summary", "Skills", "Qualification", "Certifications", "Career Summary", "Career History"]:
        section_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == label)
        assert str(section_heading.runs[0].font.color.rgb) == "000000"

    qualifications_table = next(table for table in doc.tables if table.rows[0].cells[0].text.strip() == "Qualification")
    institution_cell = qualifications_table.rows[1].cells[1]
    shd = institution_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "FFFFFF"
    for cell in qualifications_table.rows[0].cells:
        assert str(cell.paragraphs[0].runs[0].font.color.rgb) == "000000"
    for cell in qualifications_table.rows[1].cells:
        assert str(cell.paragraphs[0].runs[0].font.color.rgb) == "000000"
        assert cell.paragraphs[0].alignment == 1

    for cell in qualifications_table.rows[1].cells:
        borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
        assert borders is not None
        for edge in ("top", "left", "bottom", "right"):
            border = borders.find(qn(f"w:{edge}"))
            assert border is not None
            assert border.get(qn("w:val")) == "single"
            assert border.get(qn("w:color")) == "000000"

    career_summary_table = next(table for table in doc.tables if table.rows[0].cells[0].text.strip() == "Company")
    for cell in career_summary_table.rows[1].cells:
        assert cell.paragraphs[0].alignment == 1

    career_table = next(table for table in doc.tables if any("Company Name" in cell.text for row in table.rows for cell in row.cells))
    company_value_run = career_table.rows[0].cells[1].paragraphs[0].runs[0]
    assert company_value_run.bold is not True
    assert str(company_value_run.font.color.rgb) == "000000"
    assert round(company_value_run.font.size.pt) == 10
    assert round(career_table.rows[1].cells[1].paragraphs[0].runs[0].font.size.pt) == 10
    assert round(career_table.rows[2].cells[2].paragraphs[0].runs[0].font.size.pt) == 10
    assert round(career_table.rows[2].cells[3].paragraphs[0].runs[0].font.size.pt) == 10
    for cell in career_table.rows[1].cells:
        borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
        assert borders is not None
        for edge in ("top", "left", "bottom", "right"):
            border = borders.find(qn(f"w:{edge}"))
            assert border is not None
            assert border.get(qn("w:val")) == "single"
            assert border.get(qn("w:color")) == "000000"

    responsibility_paragraph = next(
        paragraph
        for paragraph in career_table.rows[3].cells[0].paragraphs
        if paragraph.text.strip().startswith("• ")
    )
    assert round(responsibility_paragraph.runs[0].font.size.pt) == 10


def test_profile_to_template_state_strips_leading_certification_bullets():
    profile = {
        "identity": {"full_name": "Jane Doe", "headline": "Security Analyst"},
        "summary": "Security analyst supporting governance, monitoring, and enterprise control improvements.",
        "skills": {"declared": ["SIEM"], "inferred": {}, "source_faithful": True},
        "education": [],
        "certifications": [
            "•\t• EC Council Certified Ethical Hacker CEH | EC Council",
            "•\t• Microsoft MCSA Office 365 | Microsoft",
            "•\t• IBM SOAR Specialist | IBM",
        ],
        "training": [],
        "projects": [],
        "volunteering": [],
        "publications": [],
        "languages": [],
        "awards": [],
        "interests": [],
        "references": [],
        "experience": [],
        "additional_sections": [],
        "raw_sections": [],
    }
    state = profile_to_template_state(profile)
    assert state["certifications"].splitlines() == [
        "EC Council Certified Ethical Hacker CEH | EC Council",
        "Microsoft MCSA Office 365 | Microsoft",
        "IBM SOAR Specialist | IBM",
    ]
