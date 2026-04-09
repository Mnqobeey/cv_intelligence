from io import BytesIO
from pathlib import Path
import re

from docx import Document
from fastapi.testclient import TestClient

from app.main import create_app
from app.normalizers import profile_from_sections, profile_to_template_state, validate_profile_readiness
from app.parsers import parse_experience_section, parse_sections
from app.utils_text import extract_text


BASE_DIR = Path(__file__).resolve().parents[1]
LAVINA_PDF = BASE_DIR / "uploads" / "6d7364d2-bc09-46d8-aeb4-8efdfaccd4c5.pdf"
LAVINA_DOCX = BASE_DIR / "uploads" / "6e70d4bc-9b01-4700-acd3-190bb375e723.docx"

EXPECTED_PDF_CAREER_ROWS = [
    ("Old Mutual (OMSFIN)", "Senior Project Coordinator/Project Manager", "June 2023", "June 2025"),
    ("Sanlam", "Junior Project Manager", "March 2020", "March 2022"),
    ("Interfront S.O.C", "Senior Project Administrator", "January 2019", "December 2019"),
    ("Vodacom", "Senior Project Administrator", "February 2018", "December 2018"),
    ("City of Cape Town", "Specialist Clerk - Project Administrator- Co-Ordinator", "July 2015", "December 2017"),
]

EXPECTED_PDF_CERTIFICATIONS = [
    "Scrum Master Certified | Agile Enterprise Coach | 2023",
    "Certificate - Allaboutxpert | 2013",
]


def _upload(client: TestClient, source_path: Path) -> dict:
    media_type = (
        "application/pdf"
        if source_path.suffix.lower() == ".pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response = client.post(
        "/api/upload",
        files={"file": (source_path.name, source_path.read_bytes(), media_type)},
    )
    assert response.status_code == 200, response.text
    return response.json()


def _block_by_title(blocks: list[dict], title: str) -> dict:
    for block in blocks:
        if block["title"] == title:
            return block
    raise AssertionError(f"Unable to find detected block titled {title!r}.")


def _expand_months(date_text: str) -> str:
    replacements = {
        "Jan": "January",
        "Feb": "February",
        "Mar": "March",
        "Apr": "April",
        "Jun": "June",
        "Jul": "July",
        "Aug": "August",
        "Sep": "September",
        "Oct": "October",
        "Nov": "November",
        "Dec": "December",
    }
    value = date_text
    for short, full in replacements.items():
        value = re.sub(rf"\b{short}\b", full, value)
    return value


def _validated_career_rows(validated_rows: list[dict]) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for row in validated_rows:
        normalized_dates = _expand_months(row["dates"])
        start_date, end_date = [part.strip() for part in re.split(r"\s*[–-]\s*", normalized_dates, maxsplit=1)]
        rows.append((row["company"], row["job_title"], start_date, end_date))
    return rows


def test_runtime_app_code_has_no_lavina_specific_override_markers():
    app_dir = BASE_DIR / "app"
    combined = "\n".join(path.read_text(encoding="utf-8", errors="ignore") for path in app_dir.rglob("*.py"))
    forbidden_markers = [
        "_LAVINA_",
        "_apply_lavina",
        "Lavina Jacobs targeted recovery applied",
        "benchmarked PDF/DOCX source set",
    ]
    for marker in forbidden_markers:
        assert marker not in combined


def test_lavina_pdf_upload_flow_is_source_driven_and_review_ready():
    client = TestClient(create_app())
    payload = _upload(client, LAVINA_PDF)

    state = payload["template_state"]
    workflow = payload["workflow_state"]
    summary_block = _block_by_title(payload["detected_blocks"], "Candidate Summary")
    qualification_block = _block_by_title(payload["detected_blocks"], "Qualifications")
    certification_block = _block_by_title(payload["detected_blocks"], "Certifications")

    assert state["full_name"] == "Lavina Jacobs"
    assert state["headline"] == "Project Co-ordinator"
    assert state["availability"] == "Immediately"
    assert state["region"] in {"Region not provided", "Not provided"}
    assert state["summary"].startswith("I am a highly experienced Project Coordinator")
    assert "Old Mutual" in state["summary"]
    assert "Core Skills:" not in state["skills"]
    assert "Cloud / DevOps:" not in state["skills"]
    assert "Project Resource Planning and Allocation" in state["skills"]
    assert "Azure DevOps" in state["skills"]
    assert "Project Meetings" in state["skills"]
    assert state["education"] == "Matric | Southern Suburbs Youth Academy | 2018"
    assert state["certifications"].splitlines() == EXPECTED_PDF_CERTIFICATIONS
    assert "2019 | Microsoft Digital Dashboards | Alton" in state["training"]
    assert "2023 | Managing Change | Udemy" in state["training"]
    assert "Name:" not in state["additional_sections"]
    assert "Availability:" not in state["additional_sections"]
    assert "Gender: Female" in state["additional_sections"]
    assert "Nationality: South African" in state["additional_sections"]
    assert workflow["review_ready"] is True
    assert workflow["blocking_issues"] == []
    assert payload["profile"]["document_meta"]["notes"] == []

    assert summary_block["section"] == "summary"
    assert summary_block["mapped_field"] == "summary"
    assert "highly experienced Project Coordinator" in summary_block["content"]
    assert qualification_block["mapped_field"] == "education"
    assert qualification_block["content"] == "Matric | Southern Suburbs Youth Academy | 2018"
    assert certification_block["mapped_field"] == "certifications"
    assert sorted(certification_block["content"].splitlines()) == sorted(EXPECTED_PDF_CERTIFICATIONS)
    assert not any(
        block["mapped_field"] == "education" and "highly experienced Project Coordinator" in block["content"]
        for block in payload["detected_blocks"]
    )


def test_lavina_pdf_use_suggestion_routes_summary_and_history_cleanly():
    client = TestClient(create_app())
    payload = _upload(client, LAVINA_PDF)

    summary_block = _block_by_title(payload["detected_blocks"], "Candidate Summary")
    history_block = _block_by_title(payload["detected_blocks"], "Career History")

    summary_response = client.post(
        f"/api/document/{payload['document_id']}/annotate",
        data={
            "selected_text": summary_block["content"],
            "target_key": summary_block["mapped_field"],
            "mode": "replace",
            "source_block_id": summary_block["id"],
            "source_label": summary_block["title"],
        },
    )
    assert summary_response.status_code == 200, summary_response.text
    summary_payload = summary_response.json()
    assert summary_payload["template_state"]["summary"].startswith("I am a highly experienced Project Coordinator")
    assert summary_payload["template_state"]["education"] == "Matric | Southern Suburbs Youth Academy | 2018"
    assert summary_payload["workflow_state"]["blocking_issues"] == []

    history_response = client.post(
        f"/api/document/{payload['document_id']}/annotate",
        data={
            "selected_text": history_block["content"],
            "target_key": history_block["mapped_field"],
            "mode": "replace",
            "source_block_id": history_block["id"],
            "source_label": history_block["title"],
        },
    )
    assert history_response.status_code == 200, history_response.text
    history_payload = history_response.json()
    parsed_rows = [
        (entry["company"], entry["position"], _expand_months(entry["start_date"]), _expand_months(entry["end_date"]))
        for entry in parse_experience_section(history_payload["template_state"]["career_history"])
    ]
    assert parsed_rows == EXPECTED_PDF_CAREER_ROWS
    assert history_payload["workflow_state"]["blocking_issues"] == []


def test_lavina_pdf_review_and_download_use_normalized_experience():
    client = TestClient(create_app())
    payload = _upload(client, LAVINA_PDF)

    review = client.post(
        f"/api/document/{payload['document_id']}/review-complete",
        json={"template_state": payload["template_state"]},
    )
    assert review.status_code == 200, review.text
    validated = review.json()["validated_export_json"]

    assert validated["identity"]["full_name"] == "Lavina Jacobs"
    assert validated["identity"]["headline"] == "Project Co-ordinator"
    assert validated["identity"]["availability"] == "Immediately"
    assert validated["identity"]["region"] in {"Region not provided", "Not provided"}
    assert validated["career_summary"].startswith("I am a highly experienced Project Coordinator")
    assert _validated_career_rows(validated["career_history"]) == EXPECTED_PDF_CAREER_ROWS

    download = client.post(f"/api/document/{payload['document_id']}/download", json={})
    assert download.status_code == 200, download.text

    doc = Document(BytesIO(download.content))
    paragraph_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()
    )
    full_text = paragraph_text + "\n" + table_text

    assert "Lavina Jacobs" in full_text
    assert "Project Co-ordinator" in full_text
    assert "Candidate Summary" in paragraph_text
    assert "Skills" in paragraph_text
    assert "Qualification" in paragraph_text
    assert "Certifications" in paragraph_text
    assert "Career Summary" in paragraph_text
    assert "Career History" in paragraph_text
    assert paragraph_text.index("Lavina Jacobs") < paragraph_text.index("Candidate Summary")
    assert paragraph_text.index("Candidate Summary") < paragraph_text.index("Skills")
    assert paragraph_text.index("Skills") < paragraph_text.index("Qualification")
    assert paragraph_text.index("Qualification") < paragraph_text.index("Certifications")
    assert paragraph_text.index("Certifications") < paragraph_text.index("Career Summary")
    assert paragraph_text.index("Career Summary") < paragraph_text.index("Career History")
    assert "Old Mutual (OMSFIN)" in full_text
    assert "Sanlam" in full_text
    assert "Interfront S.O.C" in full_text
    assert "Vodacom" in full_text
    assert "City of Cape Town" in full_text


def test_lavina_docx_extraction_stays_generic_and_honest_about_missing_history():
    raw_text = extract_text(LAVINA_DOCX)
    sections = parse_sections(raw_text)
    profile = profile_from_sections(raw_text, sections, LAVINA_DOCX)
    state = profile_to_template_state(profile)

    assert state["full_name"] == "Lavina Jacobs"
    assert state["headline"] == "Project Co-ordinator"
    assert state["availability"] == "Immediately"
    assert state["summary"].startswith("I am a highly experienced Project Coordinator")
    assert state["education"] == "Matric | Southern Suburbs Youth Academy"
    assert sorted(state["certifications"].splitlines()) == sorted([
        "Scrum Master Certified | Agile Enterprise Coach",
        "Certificate - Allaboutxpert",
    ])
    assert "Name:" not in state["additional_sections"]
    assert "Availability:" not in state["additional_sections"]
    assert "Gender: Female" in state["additional_sections"]
    assert "Nationality: South African" in state["additional_sections"]
    assert profile["document_meta"]["notes"] == []
    assert profile["experience"] == []
    assert validate_profile_readiness(state) == ["Career History is required before build can pass."]

    client = TestClient(create_app())
    payload = _upload(client, LAVINA_DOCX)
    summary_block = _block_by_title(payload["detected_blocks"], "Candidate Summary")
    assert summary_block["mapped_field"] == "summary"
    assert not any(
        block["mapped_field"] == "education" and "highly experienced Project Coordinator" in block["content"]
        for block in payload["detected_blocks"]
    )
    assert payload["profile"]["document_meta"]["notes"] == []
    assert payload["workflow_state"]["blocking_issues"] == ["Career History is required before build can pass."]
