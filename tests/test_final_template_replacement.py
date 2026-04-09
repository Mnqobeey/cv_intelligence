from pathlib import Path

from docx import Document

from app.constants import BASE_TEMPLATE_PATH, MASTER_TEMPLATE_PATH
from app.docx_exporter import build_profile_docx


def _state():
    return {
        "full_name": "Thandokuhle Mnqobi Mntambo",
        "headline": "QA Intern",
        "availability": "Immediate",
        "region": "Johannesburg",
        "summary": "QA-focused candidate with testing and data experience.",
        "skills": "Selenium\nPython\nPower BI",
        "education": "BSc Information Systems | Example University | 2024",
        "certifications": "Java S",
        "career_history": "CestaSoft Solutions | QA Intern | Jun 2025 | Present\nValidated releases and regression tests.",
    }


def test_final_template_asset_is_the_only_runtime_template():
    assert BASE_TEMPLATE_PATH == MASTER_TEMPLATE_PATH
    assert BASE_TEMPLATE_PATH.name == "Template Cestasoft Profile.docx"
    assets = set(p.name for p in BASE_TEMPLATE_PATH.parent.glob("*.docx"))
    assert assets == {"Template Cestasoft Profile.docx"}


def test_export_uses_replaced_final_template_layout(tmp_path: Path):
    out = tmp_path / "out.docx"
    build_profile_docx(out, _state(), profile=None)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(cell.text for t in doc.tables for r in t.rows for cell in r.cells)
    assert "Thandokuhle Mnqobi Mntambo" in text
    assert "QA Intern" in text
    assert "Candidate Summary" in text
    assert "Career Summary" in text
    assert "Career History" in text


def test_runtime_template_is_neutral_and_marker_driven(tmp_path: Path):
    out = tmp_path / "out.docx"
    build_profile_docx(out, _state(), profile=None)

    template = Document(str(MASTER_TEMPLATE_PATH))
    text = "\n".join(p.text for p in template.paragraphs) + "\n" + "\n".join(
        cell.text for table in template.tables for row in table.rows for cell in row.cells
    )

    assert "{{FULL_NAME}}" in text
    assert "{{HEADLINE}}" in text
    assert "{{AVAILABILITY_LINE}}" in text
    assert "{{REGION_LINE}}" in text
    assert "{{CANDIDATE_SUMMARY}}" in text
    assert "{{SKILLS_BLOCK}}" in text
    assert "{{QUALIFICATION_DEGREE}}" in text
    assert "{{QUALIFICATION_INSTITUTION}}" in text
    assert "{{QUALIFICATION_YEAR}}" in text
    assert "{{CERTIFICATION_ITEM}}" in text
    assert "{{CAREER_SUMMARY_COMPANY}}" in text
    assert "{{CAREER_SUMMARY_POSITION}}" in text
    assert "{{CAREER_SUMMARY_START}}" in text
    assert "{{CAREER_SUMMARY_END}}" in text
    assert "{{CAREER_ROLE_COMPANY}}" in text
    assert "{{CAREER_ROLE_TITLE}}" in text
    assert "{{CAREER_ROLE_START}}" in text
    assert "{{CAREER_ROLE_END}}" in text
    assert "{{CLIENT_ENGAGEMENTS_BLOCK}}" in text
    assert "{{PROJECT_NAME}}" in text
    assert "{{PROJECT_DETAILS}}" in text
    assert "George Thabiso Mpopo" not in text
    assert "Gijima Technologies" not in text
    assert "C# / NET" not in text
    assert "Mendix Apprentice Certificate" not in text
