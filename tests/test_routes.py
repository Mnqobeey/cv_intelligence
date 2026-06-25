from io import BytesIO
import json
from pathlib import Path
import sqlite3

from docx import Document
from fastapi.testclient import TestClient
import pytest

from app.constants import DATABASE_PATH
from app.main import create_app


def fake_openrouter_json() -> dict:
    return {
        "cestacv_version": 1,
        "identity": {
            "full_name": "Jordan Lee Carter",
            "headline": "Senior / Lead Software Engineer",
            "availability": "Immediate / Negotiable",
            "region": "South Africa",
            "email": "george@example.com",
            "phone": "",
            "location": "",
            "linkedin": "",
            "portfolio": "",
        },
        "career_summary": (
            "Experienced software engineer delivering enterprise platforms across public and private sector environments, "
            "leading modernisation work, improving delivery quality, and supporting scalable enterprise systems with disciplined execution."
        ),
        "skills": [{"category": "Core Skills", "items": ["C#", ".NET", "ASP.NET Core", "Angular", "Azure DevOps"]}],
        "qualifications": [{"qualification": "BSc Degree in Computer Science", "institution": "University of Zululand", "year": "2007"}],
        "certifications": [{"name": "TOGAF 9.2 Certification", "provider": "", "year": "2018"}],
        "training": [],
        "achievements": [],
        "languages": [],
        "interests": [],
        "references": [],
        "projects": [],
        "career_history": [
            {
                "job_title": "Senior Full-Stack Software Developer",
                "company": "Gijima Technologies",
                "start_date": "Apr 2024",
                "end_date": "Present",
                "responsibilities": ["Lead full-stack development on national-scale government systems."],
                "client_engagements": [],
                "projects": [],
            }
        ],
        "additional_sections": [],
    }


@pytest.fixture(autouse=True)
def mock_openrouter(monkeypatch):
    async def fake_structure(raw_text: str):
        assert raw_text.strip()
        return fake_openrouter_json(), {
            "provider": "openrouter",
            "model": "test/model",
            "usage": {"total_tokens": 123},
            "generation_id": "test-generation",
        }

    monkeypatch.setenv("OPENROUTER_API_KEY", "test-key")
    monkeypatch.setenv("CV_INTELLIGENCE_LLM_MODE", "required")
    monkeypatch.setattr("app.routes.structure_cv_text_with_openrouter", fake_structure)


def make_upload_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Jordan Lee Carter")
    doc.add_paragraph("Senior / Lead Software Engineer")
    doc.add_paragraph("Email: george@example.com")
    doc.add_paragraph("Region: South Africa")
    doc.add_paragraph("Candidate Summary")
    doc.add_paragraph("Experienced software engineer delivering enterprise platforms across public and private sector environments, leading modernisation work, improving delivery quality, and supporting scalable enterprise systems with disciplined execution.")
    doc.add_paragraph("Skills")
    doc.add_paragraph("C#, .NET, ASP.NET Core, Angular, Azure DevOps")
    doc.add_paragraph("Qualifications")
    doc.add_paragraph("BSc Degree in Computer Science | University of Zululand | 2007")
    doc.add_paragraph("Certifications")
    doc.add_paragraph("TOGAF 9.2 Certification (2018)")
    doc.add_paragraph("Career History")
    doc.add_paragraph("Gijima Technologies | Senior Full-Stack Software Developer | Apr 2024 | Present")
    doc.add_paragraph("Lead full-stack development on national-scale government systems.")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def upload_sample(client: TestClient):
    upload_bytes = make_upload_docx()
    response = client.post(
        "/api/upload",
        files={"file": ("candidate.docx", upload_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 200, response.text
    return response.json()


def test_upload_requires_review_before_download(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)

    download = client.post(
        f"/api/document/{data['document_id']}/download",
        json={"template_state": data["template_state"]},
    )
    assert download.status_code == 400
    assert "review" in download.json()["detail"].lower()


def test_review_completion_unlocks_download(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)

    review = client.post(
        f"/api/document/{data['document_id']}/review-complete",
        json={"template_state": data["template_state"]},
    )
    assert review.status_code == 200, review.text
    review_json = review.json()
    assert review_json["workflow_state"]["can_download"] is True
    assert "career_summary" in review_json["validated_export_json"]

    download = client.post(
        f"/api/document/{data['document_id']}/download",
        json={"template_state": data["template_state"]},
    )
    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(download.content) > 1000


def test_download_rebuilds_validated_export_when_reviewed_document_loses_cached_payload(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)

    review = client.post(
        f"/api/document/{data['document_id']}/review-complete",
        json={"template_state": data["template_state"]},
    )
    assert review.status_code == 200, review.text

    with sqlite3.connect(DATABASE_PATH) as connection:
        row = connection.execute(
            "SELECT payload_json FROM documents WHERE document_id = ?",
            (data["document_id"],),
        ).fetchone()
        payload = json.loads(row[0])
        payload.pop("validated_export_json", None)
        connection.execute(
            "UPDATE documents SET payload_json = ? WHERE document_id = ?",
            (json.dumps(payload, ensure_ascii=False), data["document_id"]),
        )
        connection.commit()

    refreshed = client.get(f"/api/document/{data['document_id']}")
    assert refreshed.status_code == 200, refreshed.text
    assert refreshed.json()["workflow_state"]["can_download"] is True

    download = client.post(
        f"/api/document/{data['document_id']}/download",
        json={"template_state": data["template_state"]},
    )
    assert download.status_code == 200, download.text
    assert len(download.content) > 1000


def test_download_revalidates_when_template_state_changes_after_review_complete(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)

    review = client.post(
        f"/api/document/{data['document_id']}/review-complete",
        json={"template_state": data["template_state"]},
    )
    assert review.status_code == 200, review.text

    mutated = dict(data["template_state"])
    mutated["summary"] = (mutated.get("summary") or "") + " UNIQUE_EXPORT_TOKEN_42"
    mutated["headline"] = "Updated Headline UNIQUE_EXPORT_HEADLINE_42"

    download = client.post(
        f"/api/document/{data['document_id']}/download",
        json={"template_state": mutated},
    )
    assert download.status_code == 200, download.text
    doc = Document(BytesIO(download.content))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    full_text += "\n" + "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    )
    assert "UNIQUE_EXPORT_TOKEN_42" in full_text
    assert "UNIQUE_EXPORT_HEADLINE_42" in full_text


def test_removed_export_endpoint_returns_410(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)

    response = client.post(f"/api/document/{data['document_id']}/export", json={})

    assert response.status_code == 410


def test_preview_is_locked_to_george_sections(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)
    preview = data["preview_html"]
    assert "Candidate Summary" in preview
    assert "Career Summary" in preview
    assert "References" not in preview
    assert "Projects" not in preview


def test_template_update_normalizes_certification_bullets_for_current_document(tmp_path: Path):
    app = create_app()
    client = TestClient(app)
    data = upload_sample(client)

    contaminated = dict(data["template_state"])
    contaminated["certifications"] = (
        "•\t• EC Council Certified Ethical Hacker CEH | EC Council\n"
        "•\t• Microsoft MCSA Office 365 | Microsoft\n"
        "•\t• IBM SOAR Specialist | IBM"
    )

    update = client.post(
        f"/api/document/{data['document_id']}/template",
        json=contaminated,
    )
    assert update.status_code == 200, update.text
    payload = update.json()
    assert payload["template_state"]["certifications"].splitlines() == [
        "EC Council Certified Ethical Hacker CEH | EC Council",
        "Microsoft MCSA Office 365 | Microsoft",
        "IBM SOAR Specialist | IBM",
    ]
    assert "•\t•" not in payload["preview_html"]
    assert "••" not in payload["preview_html"]


def test_review_complete_returns_actionable_400_when_schema_validation_fails(tmp_path: Path):
    app = create_app()
    client = TestClient(app, raise_server_exceptions=False)
    payload = {
        "cestacv_version": 1,
        "identity": {
            "full_name": "Naledi Khumalo",
            "headline": "Research Assistant",
            "availability": "",
            "region": "",
            "email": "audit@example.com",
            "phone": "0123456789",
            "location": "Johannesburg",
            "linkedin": "",
            "portfolio": "",
        },
        "career_summary": "Research assistant supporting structured delivery, documentation, and coordination across university projects and practical work environments.",
        "skills": [{"category": "Technical Skills", "items": ["Python", "SQL"]}],
        "qualifications": [{"qualification": "BSc Information Systems", "institution": "University of Johannesburg", "year": "2024"}],
        "certifications": [],
        "training": [],
        "achievements": [],
        "languages": [],
        "interests": [],
        "references": [],
        "projects": [],
        "career_history": [
            {
                "job_title": "Database Systems Module",
                "company": "University of Johannesburg",
                "start_date": "2024",
                "end_date": "2024",
                "responsibilities": ["Completed semester coursework and practical assignments"],
            }
        ],
        "additional_sections": [],
    }
    upload = client.post("/api/upload-text", json={"text": json.dumps(payload)})
    assert upload.status_code == 200, upload.text
    uploaded = upload.json()

    review = client.post(
        f"/api/document/{uploaded['document_id']}/review-complete",
        json={"template_state": uploaded["template_state"]},
    )
    assert review.status_code == 400, review.text
    detail = review.json()["detail"]
    assert detail["message"] in {
        "Profile failed schema validation.",
        "Profile is not ready for review completion.",
    }
    assert any(
        issue in detail["issues"]
        for issue in {
            "Pure education records are not allowed in career_history.",
            "Career History contains pure education records that must be removed.",
        }
    )
